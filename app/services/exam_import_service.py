"""
Service để import đề thi từ file DOCX và chuyển đổi thành JSON
"""

import logging
import json
import re
import time
import uuid
from typing import Dict, Any, Optional, List
from docx import Document
import io
from datetime import datetime

from app.services.openrouter_service import get_openrouter_service
from app.models.exam_import_models import (
    ExamImportRequest,
    ExamImportResponse,
    ExamImportError,
    ImportedExamData,
    ExamImportStatistics
)
from app.constants.difficulty_levels import DifficultyLevel

logger = logging.getLogger(__name__)


class ExamImportService:
    """Service để import và xử lý đề thi từ file DOCX"""

    def __init__(self):
        self.model_name = "google/gemini-2.0-flash-001"
        logger.info("🔄 ExamImportService: First-time initialization triggered")

    async def import_exam_from_docx_content(
        self, file_content: bytes, filename: str = "exam.docx", staff_import: bool = False
    ) -> Dict[str, Any]:
        """
        Import đề thi từ nội dung file DOCX

        Args:
            file_content: Nội dung file DOCX dưới dạng bytes
            filename: Tên file gốc
            staff_import: True nếu import cho staff (format SpringBoot), False cho frontend

        Returns:
            Dict chứa kết quả import
        """
        start_time = time.time()
        
        try:
            logger.info(f"Starting DOCX import for file: {filename}")
            
            # 1. Extract text từ DOCX
            extracted_text = self._extract_text_from_docx_bytes(file_content)
            
            if not extracted_text or len(extracted_text.strip()) < 100:
                return {
                    "statusCode": 400,
                    "message": "File extraction failed",
                    "error": "Không thể trích xuất nội dung từ file DOCX hoặc nội dung quá ngắn",
                    "details": {"filename": filename, "extracted_length": len(extracted_text)}
                }

            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")

            # 2. Validate format đề thi trước khi gọi LLM
            logger.info("Validating exam format...")
            format_validation = self._validate_exam_format(extracted_text)

            if not format_validation["is_valid"]:
                return {
                    "statusCode": 400,
                    "message": "Invalid exam format",
                    "error": f"Đề thi không đúng format chuẩn: {format_validation['error']}",
                    "details": {
                        "filename": filename,
                        "validation_details": format_validation["details"]
                    }
                }

            # Lưu thông tin warnings để trả về sau
            format_warnings = format_validation.get("warnings", [])
            missing_parts = format_validation.get("details", {}).get("missing_parts", [])

            logger.info(f"Exam format validation passed with warnings: {format_warnings}")

            # 3. Gửi cho LLM để phân tích và chuyển đổi
            logger.info("Sending content to LLM for analysis...")
            llm_result = await self._analyze_exam_with_llm(extracted_text, filename)

            if not llm_result.get("success", False):
                return {
                    "statusCode": 500,
                    "message": "LLM analysis failed",
                    "error": f"Không thể phân tích đề thi: {llm_result.get('error', 'Unknown error')}",
                    "details": {"filename": filename}
                }

            # 3. Parse JSON response từ LLM
            exam_data = llm_result.get("data")
            if not exam_data:
                return {
                    "statusCode": 500,
                    "message": "No exam data returned",
                    "error": "LLM không trả về dữ liệu đề thi",
                    "details": {"filename": filename}
                }

            # 4. Validate và clean dữ liệu từ LLM
            logger.info("Validating and cleaning LLM data...")
            logger.info(f"Exam data type: {type(exam_data)}")
            logger.info(f"Exam data preview: {str(exam_data)[:500]}...")

            validation_result = self._validate_and_clean_exam_data(exam_data)

            if not validation_result["is_valid"]:
                return {
                    "statusCode": 422,
                    "message": "Invalid exam data from LLM",
                    "error": f"Dữ liệu từ LLM không hợp lệ: {validation_result['error']}",
                    "details": {
                        "filename": filename,
                        "validation_details": validation_result["details"],
                        "exam_data_type": str(type(exam_data)),
                        "exam_data_preview": str(exam_data)[:200]
                    }
                }

            # Sử dụng dữ liệu đã được clean
            exam_data = validation_result["cleaned_data"]

            # 5. Chuyển đổi sang format phù hợp
            if staff_import:
                # Format cho SpringBoot staff
                formatted_data = self._convert_to_staff_format(exam_data)
                success_message = "Question bank data imported successfully"
            else:
                # Format cho Frontend
                formatted_data = self._convert_to_fe_format(exam_data)
                success_message = "Template updated successfully"

            # 6. Validate và tạo response
            processing_time = time.time() - start_time

            # Tạo message với thông tin về các phần thiếu
            if format_warnings:
                success_message += f" (Lưu ý: {'; '.join(format_warnings)})"

            # Tạo response theo format phù hợp
            response_data = {
                "statusCode": 200,
                "message": success_message,
                "data": formatted_data
            }

            return response_data

        except Exception as e:
            logger.error(f"Error importing exam from DOCX: {e}")
            processing_time = time.time() - start_time
            
            return {
                "statusCode": 500,
                "message": "Import failed",
                "error": f"Lỗi trong quá trình import: {str(e)}",
                "details": {
                    "filename": filename,
                    "processing_time": processing_time,
                    "error_type": type(e).__name__
                }
            }

    def _extract_text_from_docx_bytes(self, file_content: bytes) -> str:
        """
        Trích xuất text từ file DOCX bytes

        Args:
            file_content: Nội dung file DOCX dưới dạng bytes

        Returns:
            Text đã trích xuất
        """
        try:
            # Tạo file-like object từ bytes
            file_stream = io.BytesIO(file_content)
            
            # Đọc document
            doc = Document(file_stream)
            
            # Trích xuất text từ tất cả paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Trích xuất text từ tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            extracted_text = "\n".join(full_text)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            
            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""

    def _validate_exam_format(self, exam_text: str) -> Dict[str, Any]:
        """
        Validate format đề thi cơ bản - kiểm tra có ít nhất 1 phần và đáp án

        Args:
            exam_text: Nội dung đề thi đã extract

        Returns:
            Dict chứa kết quả validation
        """
        try:
            # Normalize text để dễ kiểm tra
            normalized_text = exam_text.upper().replace('\n', ' ').replace('\r', '')

            validation_result = {
                "is_valid": True,
                "error": "",
                "details": {},
                "warnings": []
            }

            # 1. Kiểm tra cấu trúc các phần
            all_parts = ["PHẦN I", "PHẦN II", "PHẦN III"]
            found_parts = []
            missing_parts = []

            for part in all_parts:
                if part in normalized_text:
                    found_parts.append(part)
                else:
                    missing_parts.append(part)

            # Chỉ yêu cầu có ít nhất 1 phần
            if not found_parts:
                validation_result["is_valid"] = False
                validation_result["error"] = "Không tìm thấy cấu trúc phần nào (PHẦN I, II, III)"
                validation_result["details"]["missing_parts"] = missing_parts
                return validation_result

            # Ghi nhận các phần thiếu như warning
            if missing_parts:
                validation_result["warnings"].append(f"Thiếu các phần: {', '.join(missing_parts)}")
                validation_result["details"]["missing_parts"] = missing_parts
                validation_result["details"]["found_parts"] = found_parts

            # 2. Bỏ qua kiểm tra phần đáp án (không bắt buộc)
            # if "ĐÁP ÁN" not in normalized_text:
            #     validation_result["warnings"].append("Không tìm thấy phần đáp án")
            #     validation_result["details"]["missing_answer_section"] = True

            logger.info(f"Exam format validation passed - Found parts: {found_parts}, Missing: {missing_parts}")
            return validation_result

        except Exception as e:
            logger.error(f"Error in format validation: {e}")
            return {
                "is_valid": False,
                "error": f"Lỗi trong quá trình validate format: {str(e)}",
                "details": {"exception": str(e)},
                "warnings": []
            }

    async def _analyze_exam_with_llm(self, exam_text: str, filename: str) -> Dict[str, Any]:
        """
        Sử dụng LLM để phân tích và chuyển đổi đề thi thành JSON

        Args:
            exam_text: Nội dung đề thi đã extract
            filename: Tên file gốc

        Returns:
            Dict chứa kết quả phân tích
        """
        try:
            # Tạo prompt cho LLM
            prompt = self._create_analysis_prompt(exam_text, filename)
            
            # Gọi LLM
            openrouter_service = get_openrouter_service()
            response = await openrouter_service.generate_content(
                prompt=prompt,
                temperature=0.1,
                max_tokens=8000
            )

            if not response.get("success", False):
                return {
                    "success": False,
                    "error": f"LLM request failed: {response.get('error', 'Unknown error')}"
                }

            # Parse JSON từ response
            llm_content = response.get("text", "") or response.get("content", "")
            logger.info(f"LLM response length: {len(llm_content)}")
            logger.info(f"LLM response preview: {llm_content[:500]}...")

            # Tìm JSON trong response
            json_data = self._extract_json_from_response(llm_content)
            
            if not json_data:
                return {
                    "success": False,
                    "error": "Không thể tìm thấy JSON hợp lệ trong response của LLM"
                }

            return {
                "success": True,
                "data": json_data
            }

        except Exception as e:
            logger.error(f"Error in LLM analysis: {e}")
            return {
                "success": False,
                "error": f"LLM analysis error: {str(e)}"
            }

    def _create_analysis_prompt(self, exam_text: str, filename: str) -> str:
        """
        Tạo prompt cho LLM để phân tích đề thi

        Args:
            exam_text: Nội dung đề thi
            filename: Tên file

        Returns:
            Prompt string
        """
        prompt = f"""
Bạn là một chuyên gia phân tích đề thi. Hãy phân tích nội dung đề thi sau và chuyển đổi thành định dạng JSON chính xác.

NỘI DUNG ĐỀ THI:
{exam_text}

YÊU CẦU:
1. Phân tích và trích xuất thông tin đề thi thành JSON với cấu trúc chính xác như mẫu
2. Xác định môn học, lớp, thời gian làm bài, tên trường:
   - subject: PHẢI xác định môn học (VD: "Hóa học", "Toán học", "Vật lý"), KHÔNG được để null
   - grade: PHẢI xác định lớp (số từ 1-12), nếu không rõ thì mặc định là 12, KHÔNG được để null
   - duration_minutes: PHẢI xác định thời gian (số phút), nếu không rõ thì mặc định 90, KHÔNG được để null
   - school: Tìm và trích xuất tên trường từ phần đầu đề thi (thường nằm dưới "BỘ GIÁO DỤC VÀ ĐÀO TẠO")
   - Ví dụ: "TRƯỜNG THPT HONG THINH" → "TRƯỜNG THPT HONG THINH"
3. Phân chia câu hỏi theo các phần có sẵn trong đề thi:
   - Phần I: Trắc nghiệm nhiều phương án lựa chọn (A, B, C, D) - nếu có
   - Phần II: Trắc nghiệm đúng/sai (a, b, c, d với true/false) - nếu có
   - Phần III: Trắc nghiệm trả lời ngắn (chỉ số) - nếu có
4. Chỉ xử lý các phần thực sự có trong đề thi, bỏ qua phần không có
5. Trích xuất đáp án chính xác từ phần đáp án (nếu có)
6. Nếu là môn Hóa học, trích xuất bảng nguyên tử khối (nếu có)

ĐỊNH DẠNG JSON MONG MUỐN:
{{
  "subject": "Hóa học",
  "grade": 12,
  "duration_minutes": 90,
  "school": "TRƯỜNG THPT HONG THINH",
  "exam_code": "1234",
  "atomic_masses": "H = 1; C = 12; N = 14; O = 16",
  "parts": [
    {{
      "part": "Phần I",
      "title": "Trắc nghiệm nhiều phương án lựa chọn",
      "description": "Mỗi câu trả lời đúng thí sinh được 0,25 điểm",
      "questions": [
        {{
          "id": 1,
          "question": "Nguyên tử carbon có bao nhiêu electron?",
          "options": {{
            "A": "4",
            "B": "6",
            "C": "8",
            "D": "12"
          }},
          "answer": "B"
        }}
      ]
    }},
    {{
      "part": "Phần II",
      "title": "Trắc nghiệm đúng/sai",
      "description": "Thí sinh trả lời từ câu 1 đến câu X. Mỗi câu có 4 phát biểu a), b), c), d)",
      "questions": [
        {{
          "id": 1,
          "question": "Cho các phát biểu về nguyên tử carbon:",
          "statements": {{
            "a": {{
              "text": "Nguyên tử carbon có 6 proton",
              "answer": true
            }},
            "b": {{
              "text": "Nguyên tử carbon có 8 neutron",
              "answer": false
            }},
            "c": {{
              "text": "Nguyên tử carbon có 6 electron",
              "answer": true
            }},
            "d": {{
              "text": "Nguyên tử carbon có khối lượng 14u",
              "answer": false
            }}
          }}
        }}
      ]
    }},
    {{
      "part": "Phần III",
      "title": "Trắc nghiệm trả lời ngắn",
      "description": "Chỉ ghi số, không ghi đơn vị. Tối đa 4 ký tự.",
      "questions": [
        {{
          "id": 1,
          "question": "Số electron trong nguyên tử carbon là bao nhiêu?",
          "answer": "6"
        }}
      ]
    }}
  ]
}}

LưU Ý QUAN TRỌNG VỀ CẤU TRÚC:
- Chỉ trả về JSON hợp lệ, không thêm text giải thích
- Đảm bảo tất cả câu hỏi và đáp án được trích xuất chính xác
- QUAN TRỌNG: KHÔNG ĐƯỢC TRẢ VỀ NULL CHO CÁC TRƯỜNG BẮT BUỘC:
  * subject: PHẢI là string không rỗng (VD: "Hóa học")
  * grade: PHẢI là số nguyên từ 1-12 (VD: 12)
  * duration_minutes: PHẢI là số nguyên dương (VD: 90)
  * Nếu không xác định được, sử dụng giá trị mặc định thay vì null
- QUAN TRỌNG: ID câu hỏi trong mỗi phần bắt đầu từ 1
  * Ví dụ: Phần I có câu 1-6, Phần II có câu 1-6, Phần III có câu 1-6
- QUAN TRỌNG: Mỗi loại câu hỏi có cấu trúc khác nhau:

  * PHẦN I (Trắc nghiệm nhiều lựa chọn): PHẢI có "options" và "answer"
    {{
      "id": 1,
      "question": "Câu hỏi cụ thể...",
      "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
      "answer": "A"
    }}

  * PHẦN II (Đúng/sai): PHẢI có "statements", KHÔNG có "answer" field
    {{
      "id": 1,
      "question": "Câu hỏi cụ thể...",
      "statements": {{
        "a": {{"text": "...", "answer": true}},
        "b": {{"text": "...", "answer": false}},
        "c": {{"text": "...", "answer": true}},
        "d": {{"text": "...", "answer": false}}
      }}
    }}

  * PHẦN III (Trả lời ngắn): PHẢI có "answer" string, KHÔNG có "options" hay "statements"
    {{
      "id": 1,
      "question": "Câu hỏi cụ thể...",
      "answer": "6"
    }}

- QUAN TRỌNG: Chỉ tạo parts cho những phần có NỘI DUNG CÂU HỎI thực tế trong đề thi
- Nếu phần đáp án có nhưng không có nội dung câu hỏi, KHÔNG tạo part cho phần đó
- Mảng "parts" có thể chứa 1, 2 hoặc 3 phần tùy theo nội dung đề thi thực tế
- Không tạo ra câu hỏi giả cho các phần không có nội dung
- Đảm bảo field "question" luôn là string không rỗng, không được null
- Ví dụ: Nếu đề thi chỉ có "PHẦN I" với nội dung câu hỏi, chỉ tạo 1 part cho Phần I, bỏ qua Phần II và III dù có trong đáp án
- QUAN TRỌNG: Giữ nguyên đáp án từ DOCX, KHÔNG được làm tròn, format hay thay đổi gì
  * Ví dụ: Nếu đáp án là "1,66" thì giữ nguyên "1,66", không làm tròn thành "2"
  * Nếu đáp án là "-1" thì giữ nguyên "-1"
  * Nếu đáp án là "27" thì giữ nguyên "27"
- QUAN TRỌNG: Trích xuất đúng tên trường từ phần đầu đề thi
  * Tìm dòng chứa tên trường (thường nằm dưới "BỘ GIÁO DỤC VÀ ĐÀO TẠO")
  * Ví dụ: "TRƯỜNG THPT ABC" → school: "TRƯỜNG THPT ABC"
  * Nếu không tìm thấy, để school: null

Hãy phân tích và trả về JSON:
"""
        return prompt

    def _extract_json_from_response(self, response_content: str) -> Optional[Dict[str, Any]]:
        """
        Trích xuất JSON từ response của LLM

        Args:
            response_content: Nội dung response từ LLM

        Returns:
            Dict chứa JSON data hoặc None nếu không tìm thấy
        """
        try:
            logger.info(f"Extracting JSON from response: {response_content[:200]}...")

            # 1. Tìm JSON trong markdown code blocks
            markdown_pattern = r'```(?:json)?\s*(\{.*?\})\s*```'
            markdown_matches = re.findall(markdown_pattern, response_content, re.DOTALL)

            for match in markdown_matches:
                try:
                    json_data = json.loads(match.strip())
                    if self._validate_exam_json_structure(json_data):
                        logger.info("Successfully extracted JSON from markdown code block")
                        return json_data
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse JSON from markdown: {e}")
                    continue

            # 2. Tìm JSON trong response (có thể có text bao quanh)
            json_pattern = r'\{.*\}'
            matches = re.findall(json_pattern, response_content, re.DOTALL)

            for match in matches:
                try:
                    # Thử parse JSON
                    json_data = json.loads(match)

                    # Validate cấu trúc cơ bản
                    if self._validate_exam_json_structure(json_data):
                        logger.info("Successfully extracted and validated JSON from LLM response")
                        return json_data

                except json.JSONDecodeError:
                    continue

            # 3. Nếu không tìm thấy JSON hợp lệ, thử parse toàn bộ response
            try:
                json_data = json.loads(response_content.strip())
                if self._validate_exam_json_structure(json_data):
                    logger.info("Successfully parsed entire response as JSON")
                    return json_data
            except json.JSONDecodeError:
                pass

            logger.error("Could not extract valid JSON from LLM response")
            logger.error(f"Response content: {response_content}")
            return None

        except Exception as e:
            logger.error(f"Error extracting JSON from response: {e}")
            return None

    def _validate_exam_json_structure(self, json_data: Dict[str, Any]) -> bool:
        """
        Validate cấu trúc JSON của đề thi - linh hoạt với parts

        Args:
            json_data: JSON data cần validate

        Returns:
            True nếu cấu trúc hợp lệ
        """
        try:
            required_fields = ["subject", "grade", "duration_minutes", "school", "parts"]

            # Kiểm tra các field bắt buộc
            for field in required_fields:
                if field not in json_data:
                    logger.error(f"Missing required field: {field}")
                    return False

            # Kiểm tra parts - cho phép empty list
            parts = json_data.get("parts", [])
            if not isinstance(parts, list):
                logger.error("Parts must be a list")
                return False

            # Nếu có parts, kiểm tra cấu trúc của từng part
            for i, part in enumerate(parts):
                if not isinstance(part, dict):
                    logger.error(f"Part {i} must be a dictionary")
                    return False

                part_required = ["part", "title", "questions"]
                for field in part_required:
                    if field not in part:
                        logger.error(f"Missing required field '{field}' in part {i}")
                        return False

                # Kiểm tra questions - cho phép empty list
                questions = part.get("questions", [])
                if not isinstance(questions, list):
                    logger.error(f"Questions in part {i} must be a list")
                    return False

            logger.info(f"JSON structure validation passed - {len(parts)} parts found")
            return True

        except Exception as e:
            logger.error(f"Error validating JSON structure: {e}")
            return False

    def _validate_and_clean_exam_data(self, exam_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate và clean dữ liệu đề thi từ LLM

        Args:
            exam_data: Dữ liệu thô từ LLM

        Returns:
            Dict chứa kết quả validation và dữ liệu đã clean
        """
        try:
            logger.info("Starting exam data validation and cleaning...")

            # Kiểm tra kiểu dữ liệu trước
            if not isinstance(exam_data, dict):
                return {
                    "is_valid": False,
                    "error": f"Exam data must be a dictionary, got {type(exam_data)}",
                    "details": {"data_type": str(type(exam_data))},
                    "cleaned_data": {}
                }

            logger.info(f"Raw exam data from LLM: {json.dumps(exam_data, ensure_ascii=False, indent=2)}")

            result = {
                "is_valid": True,
                "error": "",
                "details": {},
                "cleaned_data": {}
            }

            # 1. Validate basic fields
            required_fields = ["subject", "grade", "duration_minutes", "school", "parts"]
            for field in required_fields:
                if field not in exam_data:
                    result["is_valid"] = False
                    result["error"] = f"Missing required field: {field}"
                    return result

            # 2. Clean basic data
            # Xử lý grade - đảm bảo không bị None
            grade_value = exam_data.get("grade")
            if grade_value is None or grade_value == "":
                grade_value = 12  # Default grade
            try:
                grade_value = int(grade_value)
                if grade_value < 1 or grade_value > 12:
                    grade_value = 12
            except (ValueError, TypeError):
                grade_value = 12

            # Xử lý duration_minutes - đảm bảo không bị None
            duration_value = exam_data.get("duration_minutes")
            if duration_value is None or duration_value == "":
                duration_value = 90  # Default duration
            try:
                duration_value = int(duration_value)
                if duration_value <= 0:
                    duration_value = 90
            except (ValueError, TypeError):
                duration_value = 90

            cleaned_data = {
                "subject": str(exam_data.get("subject") or "").strip(),
                "grade": grade_value,
                "duration_minutes": duration_value,
                "school": str(exam_data.get("school") or "").strip(),
                "exam_code": str(exam_data.get("exam_code") or "").strip() if exam_data.get("exam_code") else None,
                "atomic_masses": str(exam_data.get("atomic_masses") or "").strip() if exam_data.get("atomic_masses") else None,
                "parts": []
            }

            # 3. Validate và clean parts
            parts = exam_data.get("parts", [])
            if not isinstance(parts, list):
                result["is_valid"] = False
                result["error"] = f"Parts must be a list, got {type(parts)}"
                result["details"]["parts_type"] = str(type(parts))
                return result

            cleaned_parts = []
            skipped_parts = []

            for i, part in enumerate(parts):
                try:
                    # Kiểm tra part có phải dict không
                    if not isinstance(part, dict):
                        logger.warning(f"Part {i} is not a dictionary, got {type(part)}: {part}")
                        skipped_parts.append({
                            "part_name": f"Part {i}",
                            "error": f"Part must be a dictionary, got {type(part)}"
                        })
                        continue

                    cleaned_part = self._clean_exam_part(part, i)
                    if cleaned_part["is_valid"]:
                        cleaned_parts.append(cleaned_part["data"])
                    else:
                        # Log warning và skip part này thay vì fail toàn bộ
                        part_name = part.get("part", f"Part {i}")
                        logger.warning(f"Skipping invalid part '{part_name}': {cleaned_part['error']}")
                        skipped_parts.append({
                            "part_name": part_name,
                            "error": cleaned_part["error"]
                        })
                except Exception as e:
                    logger.error(f"Error processing part {i}: {e}")
                    skipped_parts.append({
                        "part_name": f"Part {i}",
                        "error": f"Processing error: {str(e)}"
                    })

            # Chỉ fail nếu không có part nào hợp lệ
            if not cleaned_parts:
                result["is_valid"] = False
                result["error"] = "No valid parts found in exam data"
                result["details"]["skipped_parts"] = skipped_parts
                return result

            cleaned_data["parts"] = cleaned_parts
            result["cleaned_data"] = cleaned_data

            # Thêm thông tin về parts bị skip
            if skipped_parts:
                result["details"]["skipped_parts"] = skipped_parts
                logger.info(f"Exam data validation completed - {len(cleaned_parts)} valid parts, {len(skipped_parts)} skipped parts")
            else:
                logger.info(f"Exam data validation completed - {len(cleaned_parts)} parts validated")

            return result

        except Exception as e:
            logger.error(f"Error in exam data validation: {e}")
            return {
                "is_valid": False,
                "error": f"Validation error: {str(e)}",
                "details": {"exception": str(e)},
                "cleaned_data": {}
            }

    def _clean_exam_part(self, part_data: Dict[str, Any], part_index: int) -> Dict[str, Any]:
        """
        Clean và validate một phần của đề thi

        Args:
            part_data: Dữ liệu phần thô từ LLM
            part_index: Index của phần

        Returns:
            Dict chứa kết quả validation và dữ liệu đã clean
        """
        try:
            result = {
                "is_valid": True,
                "error": "",
                "data": {}
            }

            # Validate basic part fields
            required_part_fields = ["part", "title", "questions"]
            for field in required_part_fields:
                if field not in part_data:
                    result["is_valid"] = False
                    result["error"] = f"Missing field '{field}' in part"
                    return result

            # Clean part basic data
            cleaned_part = {
                "part": str(part_data.get("part", "")).strip(),
                "title": str(part_data.get("title", "")).strip(),
                "description": str(part_data.get("description", "")).strip() if part_data.get("description") else "",
                "questions": []
            }

            # Validate và clean questions
            questions = part_data.get("questions", [])
            if not isinstance(questions, list):
                result["is_valid"] = False
                result["error"] = f"Questions must be a list, got {type(questions)}"
                return result

            cleaned_questions = []
            invalid_questions = []

            for j, question in enumerate(questions):
                try:
                    # Kiểm tra question có phải dict không
                    if not isinstance(question, dict):
                        logger.warning(f"Question {j} in {cleaned_part['part']} is not a dictionary, got {type(question)}: {question}")
                        invalid_questions.append({
                            "question_index": j,
                            "error": f"Question must be a dictionary, got {type(question)}"
                        })
                        continue

                    cleaned_question = self._clean_question(question, cleaned_part["part"], j)
                    if cleaned_question["is_valid"]:
                        cleaned_questions.append(cleaned_question["data"])
                    else:
                        # Log invalid question nhưng không fail toàn bộ part
                        logger.warning(f"Skipping invalid question {j} in {cleaned_part['part']}: {cleaned_question['error']}")
                        invalid_questions.append({
                            "question_index": j,
                            "error": cleaned_question["error"]
                        })
                except Exception as e:
                    logger.error(f"Error processing question {j} in {cleaned_part['part']}: {e}")
                    invalid_questions.append({
                        "question_index": j,
                        "error": f"Processing error: {str(e)}"
                    })

            # Nếu không có câu hỏi hợp lệ nào, fail part này
            if not cleaned_questions:
                result["is_valid"] = False
                result["error"] = f"No valid questions found in part. Invalid questions: {len(invalid_questions)}"
                return result

            cleaned_part["questions"] = cleaned_questions
            result["data"] = cleaned_part

            return result

        except Exception as e:
            return {
                "is_valid": False,
                "error": f"Error cleaning part: {str(e)}",
                "data": {}
            }

    def _clean_question(self, question_data: Dict[str, Any], part_name: str, question_index: int) -> Dict[str, Any]:
        """
        Clean và validate một câu hỏi theo loại phần

        Args:
            question_data: Dữ liệu câu hỏi thô từ LLM
            part_name: Tên phần (để xác định loại câu hỏi)
            question_index: Index của câu hỏi

        Returns:
            Dict chứa kết quả validation và dữ liệu đã clean
        """
        try:
            logger.info(f"Cleaning question {question_index} in {part_name}")
            logger.info(f"Question data: {json.dumps(question_data, ensure_ascii=False, indent=2)}")

            result = {
                "is_valid": True,
                "error": "",
                "data": {}
            }

            # Validate basic question fields
            if "id" not in question_data or "question" not in question_data:
                result["is_valid"] = False
                result["error"] = "Missing 'id' or 'question' field"
                return result

            # Clean basic question data
            question_text = question_data.get("question")
            if not question_text or question_text is None or str(question_text).strip() == "":
                result["is_valid"] = False
                result["error"] = f"Question text cannot be null or empty. Got: {question_text}"
                return result

            cleaned_question = {
                "id": int(question_data.get("id", question_index + 1)),
                "question": str(question_text)
            }

            # Clean theo loại phần
            logger.info(f"Determining question type for part: '{part_name}'")

            # Xác định loại câu hỏi dựa trên tên phần
            part_name_upper = part_name.upper().strip()
            logger.info(f"Part name after processing: '{part_name_upper}'")

            # Sử dụng logic đơn giản để phân loại chính xác
            if part_name_upper == "PHẦN I":
                logger.info("Processing as MultipleChoice question (PHẦN I)")
                # MultipleChoiceQuestion
                if "options" not in question_data or "answer" not in question_data:
                    result["is_valid"] = False
                    result["error"] = "MultipleChoice question missing 'options' or 'answer'"
                    return result

                options = question_data.get("options", {})
                if not isinstance(options, dict):
                    result["is_valid"] = False
                    result["error"] = f"Options must be a dictionary, got {type(options)}"
                    return result

                # Kiểm tra có đủ options A, B, C, D không
                required_options = ["A", "B", "C", "D"]
                missing_options = [opt for opt in required_options if opt not in options]
                if missing_options:
                    result["is_valid"] = False
                    result["error"] = f"Options missing: {missing_options}"
                    return result

                cleaned_question["options"] = {
                    "A": str(options.get("A", "")),
                    "B": str(options.get("B", "")),
                    "C": str(options.get("C", "")),
                    "D": str(options.get("D", ""))
                }
                cleaned_question["answer"] = str(question_data.get("answer", ""))

            elif part_name_upper == "PHẦN II":
                logger.info("Processing as TrueFalse question (PHẦN II)")
                # TrueFalseQuestion
                if "statements" not in question_data:
                    result["is_valid"] = False
                    result["error"] = "TrueFalse question missing 'statements'"
                    return result

                statements = question_data.get("statements", {})
                if not isinstance(statements, dict):
                    result["is_valid"] = False
                    result["error"] = f"Statements must be a dictionary, got {type(statements)}"
                    return result

                # Kiểm tra có đủ statements a, b, c, d không
                required_statements = ["a", "b", "c", "d"]
                missing_statements = [stmt for stmt in required_statements if stmt not in statements]
                if missing_statements:
                    result["is_valid"] = False
                    result["error"] = f"Statements missing: {missing_statements}"
                    return result

                cleaned_statements = {}
                for key in ["a", "b", "c", "d"]:
                    stmt = statements.get(key, {})
                    if not isinstance(stmt, dict):
                        result["is_valid"] = False
                        result["error"] = f"Statement {key} must be a dictionary, got {type(stmt)}"
                        return result

                    if "text" not in stmt or "answer" not in stmt:
                        result["is_valid"] = False
                        result["error"] = f"Statement {key} missing 'text' or 'answer'"
                        return result

                    cleaned_statements[key] = {
                        "text": str(stmt.get("text", "")),
                        "answer": bool(stmt.get("answer", False))
                    }

                cleaned_question["statements"] = cleaned_statements

            elif part_name_upper == "PHẦN III":
                logger.info("Processing as ShortAnswer question (PHẦN III)")
                # ShortAnswerQuestion
                if "answer" not in question_data:
                    result["is_valid"] = False
                    result["error"] = "ShortAnswer question missing 'answer'"
                    return result

                cleaned_question["answer"] = str(question_data.get("answer", ""))

            else:
                result["is_valid"] = False
                result["error"] = f"Unknown part type: {part_name}"
                return result

            # Thêm explanation nếu có
            if "explanation" in question_data:
                cleaned_question["explanation"] = str(question_data.get("explanation", ""))

            result["data"] = cleaned_question
            return result

        except Exception as e:
            return {
                "is_valid": False,
                "error": f"Error cleaning question: {str(e)}",
                "data": {}
            }

    def _calculate_import_statistics(self, exam_data: Dict[str, Any]) -> ExamImportStatistics:
        """
        Tính toán thống kê cho đề thi đã import

        Args:
            exam_data: Dữ liệu đề thi

        Returns:
            ExamImportStatistics object
        """
        try:
            parts = exam_data.get("parts", [])
            logger.info(f"Calculating statistics for {len(parts)} parts")

            total_questions = 0
            part_1_questions = 0
            part_2_questions = 0
            part_3_questions = 0
            
            for part in parts:
                questions = part.get("questions", [])
                part_name = part.get("part", "").upper().strip()

                # Sử dụng logic so sánh chính xác như trong _clean_question
                if part_name == "PHẦN I":
                    part_1_questions = len(questions)
                    logger.info(f"PHẦN I: {len(questions)} questions")
                elif part_name == "PHẦN II":
                    part_2_questions = len(questions)
                    logger.info(f"PHẦN II: {len(questions)} questions")
                elif part_name == "PHẦN III":
                    part_3_questions = len(questions)
                    logger.info(f"PHẦN III: {len(questions)} questions")
                else:
                    logger.warning(f"Unknown part name for statistics: '{part_name}' with {len(questions)} questions")

                total_questions += len(questions)
            
            has_atomic_masses = bool(exam_data.get("atomic_masses"))
            
            # Tính chất lượng xử lý (dựa trên số câu hỏi và cấu trúc)
            processing_quality = min(1.0, (total_questions / 20) * 0.8 + 0.2)
            
            return ExamImportStatistics(
                total_questions=total_questions,
                part_1_questions=part_1_questions,
                part_2_questions=part_2_questions,
                part_3_questions=part_3_questions,
                has_atomic_masses=has_atomic_masses,
                processing_quality=processing_quality
            )

        except Exception as e:
            logger.error(f"Error calculating import statistics: {e}")
            return ExamImportStatistics(
                total_questions=0,
                part_1_questions=0,
                part_2_questions=0,
                part_3_questions=0,
                has_atomic_masses=False,
                processing_quality=0.0
            )

    def _convert_to_fe_format(self, exam_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Chuyển đổi dữ liệu exam sang format mà FE mong muốn

        Args:
            exam_data: Dữ liệu exam đã được clean

        Returns:
            Dict theo format FE
        """
        try:
            # Tạo UUID cho template

            # Chuyển đổi parts sang format FE
            fe_parts = []
            grading_config = {}

            for part in exam_data.get("parts", []):
                part_name = part.get("part", "")
                part_title = part.get("title", "")
                questions = part.get("questions", [])

                # Chuyển đổi questions với UUID và questionNumber
                fe_questions = []
                for idx, question in enumerate(questions):
                    fe_question = {
                        "id": str(uuid.uuid4()),
                        "questionNumber": idx + 1,
                        "question": question.get("question", "")
                    }

                    # Thêm fields tùy theo loại câu hỏi
                    if "options" in question and "answer" in question:
                        # Multiple choice
                        fe_question["options"] = question["options"]
                        fe_question["answer"] = question["answer"]
                    elif "statements" in question:
                        # True/False
                        fe_question["statements"] = question["statements"]
                    elif "answer" in question and "options" not in question:
                        # Short answer
                        fe_question["answer"] = question["answer"]

                    # Thêm difficultyLevel cho từng câu hỏi
                    fe_question["difficultyLevel"] = self._analyze_difficulty_level(question).value

                    fe_questions.append(fe_question)

                fe_part = {
                    "part": part_name,
                    "title": part_title,
                    "questions": fe_questions
                }

                fe_parts.append(fe_part)

                # Tạo grading config (mặc định)
                if part_name == "PHẦN I":
                    grading_config[part_name] = 0.25
                elif part_name == "PHẦN II":
                    grading_config[part_name] = 1.0
                elif part_name == "PHẦN III":
                    grading_config[part_name] = 0.25
                else:
                    grading_config[part_name] = 0.5

            # Tính tổng điểm
            total_score = 10.0

            # Tạo response theo format FE - sử dụng dữ liệu đã được clean
            fe_data = {
                "name": f"Template {exam_data.get('subject', 'Chưa xác định')}",
                "subject": exam_data.get("subject", "Chưa xác định"),
                "grade": exam_data.get("grade", "Chưa xác định"),  # Sử dụng giá trị đã clean hoặc default
                "durationMinutes": exam_data.get("duration_minutes", 90),  # Sử dụng giá trị đã clean hoặc default
                "parts": fe_parts,
                "totalScore": total_score,
                "version": 1,
                "createdAt": datetime.now().isoformat()
            }

            return fe_data

        except Exception as e:
            logger.error(f"Error converting to FE format: {e}")
            # Trả về format cơ bản nếu có lỗi
            return {
                "id": str(uuid.uuid4()),
                "name": "Template mới",
                "subject": "Chưa xác định",
                "grade": 12,
                "durationMinutes": 90,
                "createdBy": str(uuid.uuid4()),
                "contentJson": {
                    "parts": []
                },
                "gradingConfig": {},
                "totalScore": 10.0,
                "version": 1,
                "createdAt": datetime.now().isoformat()
            }

    def _convert_to_staff_format(self, exam_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chuyển đổi dữ liệu exam sang format SpringBoot cho staff

        Args:
            exam_data: Dữ liệu exam đã được clean

        Returns:
            List[Dict] theo format SpringBoot QuestionBank với lessonId = null
        """
        try:
            staff_questions = []

            for part in exam_data.get("parts", []):
                part_name = part.get("part", "")
                questions = part.get("questions", [])

                # Xác định questionType dựa trên part
                question_type = self._map_part_to_question_type(part_name)

                for question in questions:
                    # Tạo question cho SpringBoot format
                    staff_question = {
                        "lessonId": None,  # Staff tự chọn
                        "questionType": question_type,
                        "difficultyLevel": self._analyze_difficulty_level(question).value,
                        "questionContent": self._format_question_content(question, question_type),
                        "explanation": question.get("explanation", ""),
                        "referenceSource": exam_data.get("school") or "Imported from DOCX",
                        "suggest": self._generate_lesson_suggestions(question, exam_data)
                    }

                    staff_questions.append(staff_question)

            logger.info(f"Converted {len(staff_questions)} questions to staff format")
            return staff_questions

        except Exception as e:
            logger.error(f"Error converting to staff format: {e}")
            return []



    def _generate_lesson_suggestions(self, question: Dict[str, Any], exam_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Tạo gợi ý về bài học dựa trên nội dung câu hỏi

        Args:
            question: Dữ liệu câu hỏi
            exam_data: Dữ liệu đề thi

        Returns:
            Dict: Thông tin gợi ý cho staff academic
        """
        try:
            question_text = question.get("question", "").lower()

            # Phân tích keywords trong câu hỏi
            chemistry_topics = self._analyze_chemistry_topics(question_text)

            # Tạo suggestions
            suggestions = {
                "keywords": self._extract_key_concepts(question_text),
                "topics": chemistry_topics,
                "subject_name": exam_data.get("subject", ""),
                "grade_level": exam_data.get("grade", 12)
            }

            return suggestions

        except Exception as e:
            logger.error(f"Error generating lesson suggestions: {e}")
            return {
                "keywords": [],
                "topics": [],
                "subject_name": exam_data.get("subject", ""),
                "grade_level": exam_data.get("grade", 12),
                "error": "Could not generate suggestions"
            }

    def _analyze_chemistry_topics(self, question_text: str) -> List[str]:
        """
        Phân tích chủ đề hóa học từ câu hỏi

        Args:
            question_text: Nội dung câu hỏi

        Returns:
            List[str]: Danh sách chủ đề hóa học
        """
        topics = []

        # Mapping keywords to chemistry topics
        topic_keywords = {
            "Cấu tạo nguyên tử": ["nguyên tử", "proton", "neutron", "electron", "hạt nhân", "lớp vỏ", "orbital"],
            "Bảng tuần hoàn": ["bảng tuần hoàn", "chu kỳ", "nhóm", "kim loại", "phi kim", "khí hiếm"],
            "Liên kết hóa học": ["liên kết", "ion", "cộng hóa trị", "kim loại", "phân tử", "tinh thể"],
            "Phản ứng hóa học": ["phản ứng", "oxi hóa", "khử", "cân bằng", "tốc độ phản ứng"],
            "Dung dịch": ["dung dịch", "nồng độ", "mol", "độ tan", "ph", "acid", "base"],
            "Hóa hữu cơ": ["hydrocarbon", "alcohol", "acid carboxylic", "ester", "amin", "protein"],
            "Nhiệt hóa học": ["enthalpy", "entropy", "năng lượng", "nhiệt độ", "cháy"],
            "Điện hóa": ["điện phân", "pin", "thế điện cực", "ăn mòn điện hóa"]
        }

        for topic, keywords in topic_keywords.items():
            if any(keyword in question_text for keyword in keywords):
                topics.append(topic)

        return topics

    def _extract_key_concepts(self, question_text: str) -> List[str]:
        """
        Trích xuất các khái niệm chính từ câu hỏi

        Args:
            question_text: Nội dung câu hỏi

        Returns:
            List[str]: Danh sách khái niệm chính
        """
        concepts = []

        # Common chemistry concepts
        concept_patterns = [
            r'(nguyên tử \w+)', r'(phân tử \w+)', r'(ion \w+)',
            r'(axit \w+)', r'(bazơ \w+)', r'(muối \w+)',
            r'(kim loại \w+)', r'(phi kim \w+)',
            r'(phản ứng \w+)', r'(dung dịch \w+)',
            r'(nồng độ \w+)', r'(khối lượng \w+)'
        ]

        for pattern in concept_patterns:
            matches = re.findall(pattern, question_text, re.IGNORECASE)
            concepts.extend(matches)

        # Remove duplicates and limit
        return list(set(concepts))[:5]

    def _map_part_to_question_type(self, part_name: str) -> str:
        """
        Map tên phần sang questionType cho SpringBoot

        Args:
            part_name: Tên phần (Phần I, II, III)

        Returns:
            str: Question type
        """
        part_mapping = {
            "PHẦN I": "PART_I",
            "Phần I": "PART_I",
            "PHẦN II": "PART_II",
            "Phần II": "PART_II",
            "PHẦN III": "PART_III",
            "Phần III": "PART_III"
        }

        return part_mapping.get(part_name, "PART_I")

    def _analyze_difficulty_level(self, question: Dict[str, Any]) -> DifficultyLevel:
        """
        Phân tích mức độ khó của câu hỏi dựa trên nội dung
        Dựa theo cấu trúc đề thi THPT 2025: 75-80% nhận biết-thông hiểu, 20-25% vận dụng

        Args:
            question: Dữ liệu câu hỏi

        Returns:
            DifficultyLevel: Difficulty level enum (KNOWLEDGE, COMPREHENSION, APPLICATION)
        """
        try:
            question_text = question.get("question", "").lower()
            options_text = ""

            # Lấy text từ các lựa chọn nếu có
            if "options" in question:
                options = question.get("options", {})
                if isinstance(options, dict):
                    options_text = " ".join([str(v).lower() for v in options.values() if v])
                elif isinstance(options, list):
                    options_text = " ".join([str(opt).lower() for opt in options if opt])

            full_text = f"{question_text} {options_text}"

            # Keywords cho mức độ NHẬN BIẾT (Knowledge) - 40% đề thi
            knowledge_keywords = [
                # Định nghĩa, khái niệm cơ bản
                "là gì", "định nghĩa", "khái niệm", "tên gọi", "ký hiệu", "công thức phân tử",
                "công thức cấu tạo", "tên hóa học", "thuộc loại", "được gọi là",
                # Nhận biết tính chất
                "tính chất", "đặc điểm", "màu sắc", "trạng thái", "mùi", "vị",
                # Phân loại cơ bản
                "thuộc nhóm", "loại hợp chất", "phân loại", "nhóm chức",
                # Công thức và ký hiệu
                "ký hiệu hóa học", "số hiệu nguyên tử", "khối lượng nguyên tử",
                "cấu hình electron", "số electron", "số proton", "số neutron"
            ]

            # Keywords cho mức độ THÔNG HIỂU (Comprehension) - 35-40% đề thi
            comprehension_keywords = [
                # Giải thích hiện tượng
                "giải thích", "tại sao", "nguyên nhân", "do đâu", "vì sao",
                "điều kiện", "yếu tố ảnh hưởng", "cơ chế", "quá trình",
                # So sánh, phân biệt
                "so sánh", "phân biệt", "khác nhau", "giống nhau", "tương tự",
                "khác biệt", "điểm chung", "điểm khác",
                # Mối quan hệ
                "liên quan", "ảnh hưởng", "tác động", "phụ thuộc", "tỉ lệ",
                # Dự đoán tính chất
                "dự đoán", "nhận xét", "kết luận", "suy ra", "cho biết"
            ]

            # Keywords cho mức độ VẬN DỤNG (Application) - 20-25% đề thi
            application_keywords = [
                # Tính toán định lượng
                "tính", "tính toán", "xác định", "tìm", "khối lượng", "thể tích",
                "nồng độ", "số mol", "hiệu suất", "độ tan", "ph", "poh",
                "phần trăm", "tỉ lệ phần trăm", "khối lượng riêng",
                # Phân tích và đánh giá
                "phân tích", "đánh giá", "nhận định", "bình luận", "thảo luận",
                # Thiết kế thí nghiệm
                "thiết kế", "thí nghiệm", "phương pháp", "cách tiến hành",
                "quy trình", "các bước", "thực hiện",
                # Ứng dụng thực tế
                "ứng dụng", "sử dụng", "áp dụng", "trong thực tế", "trong đời sống",
                "sản xuất", "công nghiệp", "chế tạo", "điều chế"
            ]

            # Đếm điểm cho từng mức độ
            knowledge_score = sum(1 for keyword in knowledge_keywords if keyword in full_text)
            comprehension_score = sum(1 for keyword in comprehension_keywords if keyword in full_text)
            application_score = sum(1 for keyword in application_keywords if keyword in full_text)

            # Phân tích bổ sung dựa trên cấu trúc câu hỏi
            # Câu hỏi có số liệu cụ thể + đơn vị thường là vận dụng
            if re.search(r'\d+[.,]\d+|\d+\s*(g|ml|l|mol|m|%|°c)', full_text):
                application_score += 2

            # Câu hỏi có phương trình hóa học thường là thông hiểu hoặc vận dụng
            if re.search(r'[A-Z][a-z]?\s*\+|→|↔|=', full_text):
                comprehension_score += 1

            # Ưu tiên mạnh cho từ khóa thông hiểu
            if re.search(r'giải thích|so sánh|phân biệt', full_text):
                comprehension_score += 3

            # Câu hỏi có từ "tính", "tìm" thường là vận dụng
            elif re.search(r'tính|tìm', full_text):
                application_score += 2

            # Câu hỏi có "xác định" + số liệu thường là vận dụng
            elif re.search(r'xác định', full_text) and re.search(r'\d+', full_text):
                application_score += 1

            # Quyết định mức độ dựa trên điểm số
            if application_score > max(knowledge_score, comprehension_score):
                return DifficultyLevel.APPLICATION
            elif comprehension_score > knowledge_score:
                return DifficultyLevel.COMPREHENSION
            else:
                return DifficultyLevel.KNOWLEDGE

        except Exception as e:
            logger.error(f"Error analyzing difficulty level: {e}")
            return DifficultyLevel.KNOWLEDGE

    def _format_question_content(self, question: Dict[str, Any], question_type: str) -> Dict[str, Any]:
        """
        Format nội dung câu hỏi theo từng loại

        Args:
            question: Dữ liệu câu hỏi
            question_type: Loại câu hỏi (PART_I, PART_II, PART_III)

        Returns:
            Dict: Question content theo format SpringBoot
        """
        try:
            base_content = {
                "question": question.get("question", ""),
                "image": None  # Có thể thêm sau nếu cần
            }

            if question_type == "PART_I":
                # Multiple choice
                base_content.update({
                    "options": question.get("options", {}),
                    "answer": question.get("answer", "")
                })

            elif question_type == "PART_II":
                # True/False statements
                statements = question.get("statements", {})
                base_content["statements"] = statements

            elif question_type == "PART_III":
                # Short answer
                base_content["answer"] = question.get("answer", "")

            return base_content

        except Exception as e:
            logger.error(f"Error formatting question content: {e}")
            return {"question": question.get("question", ""), "image": None}


# Factory function để tạo ExamImportService instance
def get_exam_import_service() -> ExamImportService:
    """
    Tạo ExamImportService instance mới

    Returns:
        ExamImportService: Fresh instance
    """
    return ExamImportService()


