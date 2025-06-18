"""
Service để xuất giáo án ra file DOCX với format đẹp
"""

import re
import logging
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)


class DocxExportService:
    """Service để xuất giáo án ra file DOCX"""
    
    def __init__(self):
        self.output_dir = Path("exports")
        self.output_dir.mkdir(exist_ok=True)
    
    def create_lesson_plan_docx(self, lesson_plan_data: Dict[str, Any]) -> str:
        """
        Tạo file DOCX từ dữ liệu giáo án
        
        Args:
            lesson_plan_data: Dữ liệu giáo án từ API
            
        Returns:
            str: Đường dẫn file DOCX đã tạo
        """
        try:
            # Tạo document mới
            doc = Document()
            
            # Thiết lập style
            self._setup_document_styles(doc)
            
            # Lấy thông tin cơ bản
            content = lesson_plan_data.get('content', {})
            generated_plan = content.get('generated_plan', '')
            lesson_info = content.get('lesson_info', {})
            framework_used = content.get('framework_used', '')
            
            # Thêm header
            self._add_header(doc, lesson_info, framework_used)
            
            # Parse và format nội dung giáo án
            self._parse_and_format_content(doc, generated_plan)
            
            # Thêm footer
            self._add_footer(doc, lesson_plan_data.get('created_at', ''))
            
            # Tạo tên file
            lesson_id = lesson_info.get('lesson_id', 'unknown')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"giao_an_{lesson_id}_{timestamp}.docx"
            filepath = self.output_dir / filename
            
            # Lưu file
            doc.save(str(filepath))
            
            logger.info(f"Created DOCX file: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error creating DOCX: {e}")
            raise
    
    def _setup_document_styles(self, doc: Document):
        """Thiết lập styles cho document"""
        try:
            # Style cho tiêu đề chính
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Times New Roman'
            title_font.size = Pt(16)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # Style cho tiêu đề phần
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
            
            # Style cho nội dung
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.15
            
        except Exception as e:
            logger.warning(f"Could not setup custom styles: {e}")
    
    def _add_header(self, doc: Document, lesson_info: Dict[str, Any], framework_used: str):
        """Thêm header cho document"""
        # Tiêu đề chính
        title = lesson_info.get('lesson_title', 'GIÁO ÁN')
        title_para = doc.add_paragraph(title.upper())
        try:
            title_para.style = 'CustomTitle'
        except:
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.bold = True
                run.font.size = Pt(16)
        
        # Thông tin framework
        if framework_used:
            framework_para = doc.add_paragraph(f"(Theo {framework_used})")
            framework_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in framework_para.runs:
                run.font.italic = True
        
        # Thêm dòng trống
        doc.add_paragraph()
    
    def _parse_and_format_content(self, doc: Document, content: str):
        """Parse và format nội dung giáo án"""
        lines = content.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                # Thêm dòng trống nhỏ
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(3)
                continue

            # Kiểm tra các loại heading
            if self._is_main_heading(line):
                self._add_main_heading(doc, line)
            elif self._is_sub_heading(line):
                self._add_sub_heading(doc, line)
            elif self._is_activity_heading(line):
                self._add_activity_heading(doc, line)
            elif line.startswith('*') or line.startswith('-'):
                self._add_bullet_point(doc, line)
            elif self._is_numbered_list(line):
                self._add_numbered_list(doc, line)
            else:
                self._add_normal_paragraph(doc, line)
    
    def _is_main_heading(self, line: str) -> bool:
        """Kiểm tra có phải heading chính không"""
        patterns = [
            r'^\*\*[IVX]+\.\s*[^*]+\*\*$',  # **I. HEADING**
            r'^\*\*[A-Z][^*]*\*\*$',       # **HEADING**
            r'^[IVX]+\.\s*[A-Z]',          # I. HEADING
        ]
        return any(re.match(pattern, line) for pattern in patterns)
    
    def _is_sub_heading(self, line: str) -> bool:
        """Kiểm tra có phải sub heading không"""
        patterns = [
            r'^\*\*[0-9]+\.\s*[^*]+\*\*$',  # **1. Sub heading**
            r'^\*\*[a-z]\)\s*[^*]+\*\*$',   # **a) Sub heading**
            r'^[0-9]+\.\s*[A-Z]',           # 1. Sub heading
        ]
        return any(re.match(pattern, line) for pattern in patterns)
    
    def _is_activity_heading(self, line: str) -> bool:
        """Kiểm tra có phải activity heading không"""
        patterns = [
            r'^\*\*\([IVX]+\)\s*Hoạt động',  # **(I) Hoạt động**
            r'^\([IVX]+\)\s*Hoạt động',      # (I) Hoạt động
            r'^\*\*Hoạt động\s*[0-9]+',      # **Hoạt động 1**
        ]
        return any(re.match(pattern, line) for pattern in patterns)
    
    def _add_main_heading(self, doc: Document, line: str):
        """Thêm heading chính"""
        # Loại bỏ markdown formatting
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        para = doc.add_paragraph(clean_line)
        try:
            para.style = 'CustomHeading'
        except:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'

        # Thêm spacing
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.keep_with_next = True
    
    def _add_sub_heading(self, doc: Document, line: str):
        """Thêm sub heading"""
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        para = doc.add_paragraph(clean_line)
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    
    def _add_activity_heading(self, doc: Document, line: str):
        """Thêm activity heading"""
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        para = doc.add_paragraph(clean_line)
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(12)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
    
    def _add_bullet_point(self, doc: Document, line: str):
        """Thêm bullet point"""
        # Loại bỏ ký tự bullet
        clean_line = re.sub(r'^[\*\-]\s*', '', line)
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_line)  # Bold text
        
        para = doc.add_paragraph()
        para.style = 'List Bullet'
        para.add_run(clean_line)
        para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_numbered_list(self, doc: Document, line: str):
        """Thêm numbered list"""
        # Kiểm tra pattern như "1.", "2.", "a)", "b)"
        clean_line = re.sub(r'^[0-9]+\.\s*|^[a-z]\)\s*', '', line)
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_line)

        para = doc.add_paragraph()
        para.style = 'List Number'
        para.add_run(clean_line)
        para.paragraph_format.left_indent = Inches(0.5)

    def _is_numbered_list(self, line: str) -> bool:
        """Kiểm tra có phải numbered list không"""
        patterns = [
            r'^[0-9]+\.\s+',  # 1. 2. 3.
            r'^[a-z]\)\s+',   # a) b) c)
        ]
        return any(re.match(pattern, line) for pattern in patterns)

    def _add_normal_paragraph(self, doc: Document, line: str):
        """Thêm paragraph thường"""
        # Xử lý markdown formatting
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)

        para = doc.add_paragraph()

        # Xử lý text có bold
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.font.bold = True
            else:
                # Normal text
                para.add_run(part)

        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_footer(self, doc: Document, created_at: str):
        """Thêm footer"""
        doc.add_paragraph()  # Dòng trống
        
        # Thông tin tạo file
        footer_text = f"Giáo án được tạo tự động vào {created_at}"
        footer_para = doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.italic = True
            run.font.size = Pt(10)


# Global instance
docx_export_service = DocxExportService()
