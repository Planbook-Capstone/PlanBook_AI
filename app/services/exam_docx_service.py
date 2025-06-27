"""
Service để xuất đề thi ra file DOCX
"""

import logging
import os
import re
import tempfile
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
WD_PARAGRAPH_ALIGNMENT = WD_ALIGN_PARAGRAPH

from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from datetime import datetime

logger = logging.getLogger(__name__)


class ExamDocxService:
    """Service để tạo file DOCX cho đề thi"""

    def __init__(self):
        # Sử dụng thư mục tạm thời của hệ thống
        self.temp_dir = tempfile.gettempdir()

    async def create_exam_docx(
        self, exam_data: Dict[str, Any], exam_request: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Tạo file DOCX cho đề thi

        Args:
            exam_data: Dữ liệu đề thi đã tạo
            exam_request: Thông tin request gốc

        Returns:
            Dict chứa thông tin file đã tạo
        """
        try:
            # Tạo document mới
            doc = Document()

            # Thiết lập margins và font
            self._setup_document_style(doc)

            # Tạo header đề thi
            self._create_exam_header(doc, exam_request, exam_data)

            # Tạo thông tin đề thi
            self._create_exam_info(doc, exam_request, exam_data)

            # Tạo bảng hóa trị cho môn Hóa học
            mon_hoc = exam_request.get("mon_hoc", "").lower()
            logger.info(f"Subject check: '{exam_request.get('mon_hoc', '')}' -> '{mon_hoc}'")
            if "hóa" in mon_hoc or "chemistry" in mon_hoc:
                logger.info("Creating chemistry valence table...")
                self._create_chemistry_valence_table(doc, exam_data.get("questions", []))
            else:
                logger.info(f"Skipping chemistry valence table for subject: {mon_hoc}")

            # Tạo câu hỏi
            self._create_questions_section(doc, exam_data.get("questions", []))

            # Thêm chữ "Hết" sau phần câu hỏi, trước phần đáp án
            self._add_exam_ending(doc)

            # Tạo đáp án theo cấu trúc THPT (cùng trang với đề thi)
            self._create_thpt_answer_table(doc, exam_data.get("questions", []))

            # Lưu file với filename an toàn (không có ký tự đặc biệt) trong thư mục tạm thời
            exam_id_safe = self._sanitize_filename(
                exam_request.get("exam_id", "unknown")
            )
            filename = (
                f"exam_{exam_id_safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            filepath = os.path.join(self.temp_dir, filename)

            logger.info(f"Saving DOCX file to: {filepath}")
            logger.info(f"Temp directory: {self.temp_dir}")
            logger.info(f"Filename: {filename}")
            logger.info(f"Directory exists: {os.path.exists(self.temp_dir)}")
            logger.info(f"Directory writable: {os.access(self.temp_dir, os.W_OK)}")

            doc.save(filepath)

            return {
                "success": True,
                "filename": filename,
                "filepath": filepath,
                "file_size": os.path.getsize(filepath),
                "created_at": datetime.now().isoformat(),
            }

        except Exception as e:
            logger.error(f"Error creating exam DOCX: {e}")
            return {"success": False, "error": str(e)}

    def _sanitize_filename(self, filename: str) -> str:
        """
        Làm sạch filename để tránh lỗi encoding và giới hạn độ dài

        Args:
            filename: Tên file gốc

        Returns:
            Tên file đã được làm sạch (chỉ chứa ASCII, độ dài hợp lệ)
        """
        try:
            # Chuyển đổi ký tự tiếng Việt sang ASCII
            import unicodedata

            # Normalize unicode và loại bỏ dấu
            normalized = unicodedata.normalize('NFD', filename)
            ascii_only = normalized.encode('ascii', 'ignore').decode('ascii')

            # Loại bỏ ký tự đặc biệt, chỉ giữ lại chữ cái, số, dấu gạch dưới và gạch ngang
            sanitized = re.sub(r"[^\w\-_]", "_", ascii_only)

            # Loại bỏ nhiều dấu gạch dưới liên tiếp
            sanitized = re.sub(r"_+", "_", sanitized)

            # Loại bỏ dấu gạch dưới ở đầu và cuối
            sanitized = sanitized.strip("_")

            # Giới hạn độ dài tên file (Windows có giới hạn 255 ký tự cho tên file)
            # Để an toàn, giới hạn ở 100 ký tự
            if len(sanitized) > 100:
                sanitized = sanitized[:100]

            # Nếu kết quả rỗng, dùng default
            if not sanitized:
                sanitized = "lesson"

            return sanitized

        except Exception as e:
            logger.warning(f"Error sanitizing filename '{filename}': {e}")
            return "lesson"

    def _setup_document_style(self, doc: Document):
        """Thiết lập style cho document"""
        try:
            # Thiết lập margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

                # Bỏ số trang theo yêu cầu
                # footer = section.footer
                # footer_para = footer.paragraphs[0]
                # footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                pass

            # Thiết lập font mặc định
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

        except Exception as e:
            logger.error(f"Error setting up document style: {e}")

    def _create_exam_header(
        self, doc: Document, exam_request: Dict[str, Any], exam_data: Dict[str, Any]
    ):
        """Tạo header cho đề thi theo chuẩn THPT"""
        try:
            # Tạo bảng header 2 cột
            header_table = doc.add_table(rows=1, cols=2)
            header_table.autofit = False
            header_table.columns[0].width = Inches(3.0)
            header_table.columns[1].width = Inches(3.5)

            # Cột trái - Thông tin trường/sở
            left_cell = header_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Logo/Tên cơ quan
            left_para.add_run("BỘ GIÁO DỤC VÀ ĐÀO TạO\n").bold = True
            school_name = exam_request.get('ten_truong', 'TRƯỜNG ...')
            left_para.add_run(f"{school_name}\n").bold = True
            # Bỏ text "(Đề có ... trang)" theo yêu cầu

            # Cột phải - Thông tin đề thi
            right_cell = header_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Loại đề thi
            exam_type = "ĐỀ KIỂM TRA"
            if exam_request.get('tong_so_cau', 0) >= 40:
                exam_type = "ĐỀ THI"

            right_para.add_run(f"{exam_type} LỚP {exam_request.get('lop', ...)}\n").bold = True
            right_para.add_run(f"Môn: {exam_request.get('mon_hoc', '').upper()}\n").bold = True

            # Thời gian làm bài
            total_questions = exam_request.get('tong_so_cau', 0)
            time_minutes = 45 if total_questions <= 20 else 50 if total_questions <= 30 else 90
            right_para.add_run(f"Thời gian làm bài: {time_minutes} phút, không kể thời gian phát đề")

            # Xóa border của bảng
            for row in header_table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        OxmlElement('w:tcBorders')
                    )

            # Khoảng trống
            doc.add_paragraph()

            # Đường kẻ ngang (nét liền không có khoảng cách)
            separator = doc.add_paragraph()
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Căn phải
            separator.add_run("_" * 80)  # Nét liền bằng dấu gạch dưới

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating exam header: {e}")

    def _create_exam_info(
        self, doc: Document, exam_request: Dict[str, Any], exam_data: Dict[str, Any]
    ):
        """Tạo thông tin đề thi theo chuẩn THPT"""
        try:
            # Thông tin học sinh theo format chuẩn THPT
            info_para = doc.add_paragraph()
            info_para.add_run("Họ, tên thí sinh: ").bold = True
            info_para.add_run("." * 70)

            # Số báo danh
            sbd_para = doc.add_paragraph()
            sbd_para.add_run("Số báo danh: ").bold = True
            sbd_para.add_run("." * 70)

            # Thống kê đề thi theo chuẩn THPT
            stats = exam_data.get("statistics", {})
            total_questions = stats.get('tong_so_cau', 0)

            # Phần I - Câu trắc nghiệm nhiều phương án lựa chọn
            tn_count = stats.get("TN", 0)
            if tn_count > 0:
                part1_para = doc.add_paragraph()
                part1_run = part1_para.add_run(
                    f"PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. "
                    f"Thí sinh trả lời từ câu 1 đến câu {tn_count}. "
                    f"Mỗi câu hỏi thí sinh chỉ chọn một phương án."
                )
                part1_run.font.name = "Times New Roman"
                part1_run.font.size = Pt(11)
                part1_run.bold = True

            # Phần II - Câu đúng/sai (nếu có)
            ds_count = stats.get("DS", 0)
            if ds_count > 0:
                part2_para = doc.add_paragraph()
                start_num = tn_count + 1
                end_num = tn_count + ds_count
                part2_run = part2_para.add_run(
                    f"PHẦN II. Câu trắc nghiệm đúng sai. "
                    f"Thí sinh trả lời từ câu {start_num} đến câu {end_num}. "
                    f"Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai."
                )
                part2_run.font.name = "Times New Roman"
                part2_run.font.size = Pt(11)
                part2_run.bold = True

            # Phần III - Câu trả lời ngắn/tự luận (nếu có)
            tl_count = stats.get("TL", 0)
            dt_count = stats.get("DT", 0)
            part3_count = tl_count + dt_count
            if part3_count > 0:
                part3_para = doc.add_paragraph()
                start_num = tn_count + ds_count + 1
                end_num = total_questions
                part3_run = part3_para.add_run(
                    f"PHẦN III. Câu trả lời ngắn. "
                    f"Thí sinh trả lời từ câu {start_num} đến câu {end_num}."
                )
                part3_run.font.name = "Times New Roman"
                part3_run.font.size = Pt(11)
                part3_run.bold = True

            # doc.add_paragraph()  # Khoảng trống

     


        except Exception as e:
            logger.error(f"Error creating exam info: {e}")

    def _extract_atomic_masses_from_questions(self, questions: Optional[List[Dict[str, Any]]]) -> str:
        """
        Extract nguyên tử khối từ nội dung câu hỏi

        Args:
            questions: Danh sách câu hỏi

        Returns:
            String chứa thông tin nguyên tử khối theo format chuẩn
        """
        try:
            if not questions:
                # Default atomic masses cho Hóa học 12
                return "H = 1, C = 12, N = 14, O = 16, Al = 27, S = 32, K = 39, Fe = 56"

            # Tập hợp các nguyên tố đã tìm thấy
            found_elements = {}

            # Danh sách nguyên tố phổ biến và nguyên tử khối
            common_elements = {
                'H': 1, 'He': 4, 'Li': 7, 'Be': 9, 'B': 11, 'C': 12, 'N': 14, 'O': 16,
                'F': 19, 'Ne': 20, 'Na': 23, 'Mg': 24, 'Al': 27, 'Si': 28, 'P': 31, 'S': 32,
                'Cl': 35.5, 'Ar': 40, 'K': 39, 'Ca': 40, 'Sc': 45, 'Ti': 48, 'V': 51, 'Cr': 52,
                'Mn': 55, 'Fe': 56, 'Co': 59, 'Ni': 59, 'Cu': 64, 'Zn': 65, 'Ga': 70, 'Ge': 73,
                'As': 75, 'Se': 79, 'Br': 80, 'Kr': 84, 'Rb': 85, 'Sr': 88, 'Y': 89, 'Zr': 91,
                'Nb': 93, 'Mo': 96, 'Tc': 98, 'Ru': 101, 'Rh': 103, 'Pd': 106, 'Ag': 108, 'Cd': 112,
                'In': 115, 'Sn': 119, 'Sb': 122, 'Te': 128, 'I': 127, 'Xe': 131, 'Cs': 133, 'Ba': 137
            }

            # Tìm kiếm nguyên tố trong nội dung câu hỏi
            import re
            for question in questions:
                # Lấy nội dung từ các trường có thể chứa text
                content_fields = ['cau_hoi', 'noi_dung', 'de_bai', 'giai_thich']
                full_content = ""

                for field in content_fields:
                    if field in question and question[field]:
                        full_content += str(question[field]) + " "

                # Tìm các ký hiệu nguyên tố (1-2 chữ cái, chữ đầu viết hoa)
                element_pattern = r'\b([A-Z][a-z]?)\b'
                matches = re.findall(element_pattern, full_content)

                for element in matches:
                    if element in common_elements:
                        found_elements[element] = common_elements[element]

            # Nếu không tìm thấy nguyên tố nào, sử dụng default
            if not found_elements:
                return "H = 1, C = 12, N = 14, O = 16, Al = 27, S = 32, K = 39, Fe = 56"

            # Sắp xếp theo thứ tự alphabet và tạo string
            sorted_elements = sorted(found_elements.items())
            atomic_mass_parts = []

            for element, mass in sorted_elements:
                # Format số nguyên hoặc thập phân
                if isinstance(mass, float) and mass.is_integer():
                    atomic_mass_parts.append(f"{element} = {int(mass)}")
                else:
                    atomic_mass_parts.append(f"{element} = {mass}")

            result = ", ".join(atomic_mass_parts)

            # Đảm bảo có ít nhất một số nguyên tố cơ bản
            essential_elements = ['H', 'C', 'N', 'O']
            missing_essentials = []

            for essential in essential_elements:
                if essential not in found_elements:
                    missing_essentials.append(f"{essential} = {common_elements[essential]}")

            if missing_essentials:
                if result:
                    result = ", ".join(missing_essentials) + ", " + result
                else:
                    result = ", ".join(missing_essentials)

            return result

        except Exception as e:
            logger.error(f"Error extracting atomic masses: {e}")
            # Fallback to default
            return "H = 1, C = 12, N = 14, O = 16, Al = 27, S = 32, K = 39, Fe = 56"

    def _create_chemistry_valence_table(self, doc: Document, questions: Optional[List[Dict[str, Any]]] = None):
        """Tạo bảng hóa trị cho môn Hóa học"""
        try:
            # Extract nguyên tử khối từ questions hoặc sử dụng default
            atomic_masses = self._extract_atomic_masses_from_questions(questions)

            # Thêm bảng hóa trị cho môn Hóa học
            valence_para = doc.add_paragraph()
            valence_para.add_run(f"Cho biết nguyên tử khối: {atomic_masses}").font.size = Pt(12)
            valence_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Thêm khoảng trống
            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating chemistry valence table: {e}")

    def _add_exam_ending(self, doc: Document):
        """Thêm chữ 'Hết' ở cuối đề thi"""
        try:
            # Thêm khoảng trống
            doc.add_paragraph()

            # Thêm chữ "Hết" căn giữa
            ending_para = doc.add_paragraph()
            ending_run = ending_para.add_run("--- Hết ---")
            ending_run.bold = True
            ending_run.font.size = Pt(12)
            ending_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            logger.error(f"Error adding exam ending: {e}")

    def _create_questions_section(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo phần câu hỏi"""
        try:
            # Tiêu đề phần câu hỏi
            questions_title = doc.add_heading("PHẦN I: CÂU HỎI", level=2)
            questions_title_run = questions_title.runs[0]
            questions_title_run.font.name = "Times New Roman"
            questions_title_run.font.size = Pt(14)

            # Tạo từng câu hỏi
            for i, question in enumerate(questions, 1):
                self._create_single_question(doc, question, i)

        except Exception as e:
            logger.error(f"Error creating questions section: {e}")

    def _create_single_question(
        self, doc: Document, question: Dict[str, Any], question_num: int
    ):
        """Tạo một câu hỏi"""
        try:
            loai_cau = question.get("loai_cau", "")
            dap_an = question.get("dap_an", {})

            # Debug log để kiểm tra dữ liệu câu hỏi
            logger.info(f"Question {question_num}: loai_cau='{loai_cau}', dap_an keys: {list(dap_an.keys()) if dap_an else 'None'}")

            # Số thứ tự và nội dung câu hỏi
            q_para = doc.add_paragraph()
            q_para.add_run(f"Câu {question_num}: ").bold = True
            q_para.add_run(question.get("noi_dung_cau_hoi", ""))

            # Tạo đáp án theo loại câu hỏi
            if loai_cau == "TN":
                logger.info(f"Creating TN (multiple choice) answers for question {question_num}")
                self._create_multiple_choice_answers(doc, dap_an)
            elif loai_cau == "DT":
                logger.info(f"Creating DT (fill blank) answers for question {question_num}")
                self._create_fill_blank_answer(doc)
            elif loai_cau == "DS":
                logger.info(f"Creating DS (true/false) answers for question {question_num}")
                self._create_true_false_answers(doc, dap_an)
            elif loai_cau == "TL":
                logger.info(f"Creating TL (essay) answers for question {question_num}")
                self._create_essay_answer_space(doc)
            else:
                # Mặc định là trắc nghiệm nếu không xác định được loại
                logger.info(f"Creating default multiple choice answers for question {question_num} (loai_cau='{loai_cau}')")
                self._create_multiple_choice_answers(doc, dap_an)

            doc.add_paragraph()  # Khoảng trống giữa các câu

        except Exception as e:
            logger.error(f"Error creating single question: {e}")

    def _create_multiple_choice_answers(self, doc: Document, dap_an: Dict[str, Any]):
        """Tạo đáp án trắc nghiệm theo hàng ngang không viền"""
        try:
            options = ["A", "B", "C", "D"]
            available_options = []

            # Debug log để kiểm tra dữ liệu
            logger.info(f"Creating multiple choice answers with data: {dap_an}")

            # Lấy các đáp án có sẵn
            for option in options:
                if option in dap_an:
                    available_options.append((option, dap_an[option]))

            # Nếu không có đáp án, báo lỗi
            if not available_options:
                logger.error(f"ERROR: No answer options (A, B, C, D) found in dap_an: {dap_an}")
                logger.error(f"Available keys in dap_an: {list(dap_an.keys()) if dap_an else 'None'}")
                raise ValueError(f"Missing answer options A, B, C, D in question data: {dap_an}")
                return

            # Tạo bảng 2x2 cho 4 đáp án (2 hàng, 2 cột) KHÔNG VIỀN
            if len(available_options) <= 4:
                table = doc.add_table(rows=2, cols=2)

                # Bỏ viền bảng - không set style để tránh lỗi
                # table.style = 'Table Normal'
                for row in table.rows:
                    for cell in row.cells:
                        # Bỏ tất cả viền
                        cell._element.get_or_add_tcPr().append(
                            parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>')
                        )

                # Thiết lập độ rộng cột
                for col in table.columns:
                    col.width = Inches(3.0)

                # Điền đáp án vào bảng
                for i, (option, text) in enumerate(available_options):
                    row_idx = i // 2  # 0 hoặc 1
                    col_idx = i % 2   # 0 hoặ 1

                    if row_idx < 2 and col_idx < 2:
                        cell = table.cell(row_idx, col_idx)
                        cell_para = cell.paragraphs[0]
                        cell_para.add_run(f"{option}. ").bold = True
                        cell_para.add_run(text)

                        # Thiết lập font
                        for run in cell_para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

            # Thêm khoảng trống sau đáp án
            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating multiple choice answers: {e}")

    def _create_fill_blank_answer(self, doc: Document):
        """Tạo chỗ trống cho câu điền từ"""
        try:
            answer_para = doc.add_paragraph()
            answer_para.add_run("   Đáp án: ")
            answer_para.add_run("_" * 30)

        except Exception as e:
            logger.error(f"Error creating fill blank answer: {e}")

    def _create_true_false_answers(self, doc: Document, dap_an: Dict[str, Any]):
        """Tạo đáp án đúng/sai"""
        try:
            options = ["a", "b", "c", "d"]
            for option in options:
                option_para = doc.add_paragraph()
                option_para.add_run(f"   {option}) ").bold = True
                option_para.add_run("Đúng ☐    Sai ☐")

        except Exception as e:
            logger.error(f"Error creating true/false answers: {e}")

    def _create_essay_answer_space(self, doc: Document):
        """Tạo chỗ trống cho câu tự luận"""
        try:
            # Thêm dòng kẻ cho học sinh viết
            for i in range(5):
                line_para = doc.add_paragraph()
                line_para.add_run("_" * 80)

        except Exception as e:
            logger.error(f"Error creating essay answer space: {e}")

    def _create_thpt_answer_table(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo bảng đáp án theo cấu trúc THPT chuẩn - động theo số câu thực tế"""
        try:
            # Bỏ ngắt trang - đáp án cùng trang với đề thi
            # doc.add_page_break()

            # Thêm khoảng trống trước đáp án
            doc.add_paragraph()
            doc.add_paragraph()

            # Tiêu đề đáp án
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run("ĐÁP ÁN")
            title_run.bold = True
            title_run.font.size = Pt(14)

            doc.add_paragraph()  # Dòng trống

            # Phân loại câu hỏi theo loại
            tn_questions = [q for q in questions if q.get("loai_cau") == "TN"]
            ds_questions = [q for q in questions if q.get("loai_cau") == "DS"]
            dt_questions = [q for q in questions if q.get("loai_cau") == "DT"]

            # PHẦN I: Câu trắc nghiệm nhiều phương án lựa chọn
            if tn_questions:
                self._create_tn_answer_section(doc, tn_questions)

            # PHẦN II: Câu trắc nghiệm đúng sai
            if ds_questions:
                self._create_ds_answer_section(doc, ds_questions)

            # PHẦN III: Câu trắc nghiệm trả lời ngắn
            if dt_questions:
                self._create_dt_answer_section(doc, dt_questions)

        except Exception as e:
            logger.error(f"Error creating THPT answer table: {e}")

    def _create_tn_answer_section(self, doc: Document, tn_questions: List[Dict[str, Any]]):
        """Tạo phần đáp án trắc nghiệm nhiều phương án lựa chọn - động theo số câu thực tế"""
        try:
            # Số câu thực tế (không cố định 60)
            total_questions = len(tn_questions)

            if total_questions == 0:
                return

            # Tiêu đề phần
            section1_para = doc.add_paragraph()
            section1_run = section1_para.add_run("PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. ")
            section1_run.bold = True
            section1_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {total_questions}")

            note_para = doc.add_paragraph()
            note_para.add_run("(Mỗi câu trả lời đúng thí sinh được 0,25 điểm)")

            doc.add_paragraph()  # Dòng trống

            # Tính toán số hàng và cột dựa trên số câu thực tế
            cols_per_row = 10  # 10 cột (1 cột label + 9 cột câu hỏi)
            questions_per_row = cols_per_row - 1  # 9 câu hỏi mỗi hàng
            num_rows = (total_questions + questions_per_row - 1) // questions_per_row  # Làm tròn lên

            # Tạo bảng với số hàng động (mỗi hàng có 2 dòng: câu + đáp án)
            table = doc.add_table(rows=num_rows * 2, cols=cols_per_row)
            table.style = 'Table Grid'

            # Thiết lập độ rộng cột đều nhau
            for col in table.columns:
                col.width = Inches(0.6)

            # Điền dữ liệu vào bảng
            for row_idx in range(num_rows):
                # Header row (số câu)
                header_row = row_idx * 2
                answer_row = row_idx * 2 + 1

                # Đặt header cho hàng
                table.cell(header_row, 0).text = "Câu"
                table.cell(header_row, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(header_row, 0).paragraphs[0].runs[0].bold = True

                table.cell(answer_row, 0).text = "Chọn"
                table.cell(answer_row, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(answer_row, 0).paragraphs[0].runs[0].bold = True

                # Điền số câu và đáp án cho các cột (bỏ cột đầu)
                for col_idx in range(1, cols_per_row):
                    # Tính chỉ số câu hỏi: mỗi hàng có questions_per_row câu (bỏ cột đầu)
                    question_idx = row_idx * questions_per_row + col_idx - 1

                    # Chỉ tạo ô cho số câu thực tế
                    if question_idx < total_questions:
                        # Số câu
                        table.cell(header_row, col_idx).text = str(question_idx + 1)
                        table.cell(header_row, col_idx).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        table.cell(header_row, col_idx).paragraphs[0].runs[0].bold = True

                        # Đáp án
                        dap_an = tn_questions[question_idx].get("dap_an", {})
                        correct_answer = dap_an.get("dung", "")
                        if not correct_answer:
                            giai_thich = tn_questions[question_idx].get("giai_thich", "")
                            correct_answer = self._extract_correct_answer_from_explanation(giai_thich, dap_an)
                        if not correct_answer:
                            correct_answer = "A"  # Default nếu không tìm được

                        table.cell(answer_row, col_idx).text = correct_answer
                        table.cell(answer_row, col_idx).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if correct_answer:
                            table.cell(answer_row, col_idx).paragraphs[0].runs[0].bold = True
                    else:
                        # Ô ngoài số câu thực tế - để trống
                        table.cell(header_row, col_idx).text = ""
                        table.cell(answer_row, col_idx).text = ""

        except Exception as e:
            logger.error(f"Error creating TN answer section: {e}")

    def _create_ds_answer_section(self, doc: Document, ds_questions: List[Dict[str, Any]]):
        """Tạo phần đáp án trắc nghiệm đúng sai"""
        try:
            doc.add_paragraph()  # Dòng trống

            # Tiêu đề phần
            section2_para = doc.add_paragraph()
            section2_run = section2_para.add_run("PHẦN II. Câu trắc nghiệm đúng sai. ")
            section2_run.bold = True
            section2_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(ds_questions)}. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.")

            # Hướng dẫn chấm điểm
            note_para = doc.add_paragraph()
            note_para.add_run("- Thí sinh chỉ lựa chọn chính xác 01 ý trong 01 câu hỏi được 0,1 điểm;")
            doc.add_paragraph().add_run("- Thí sinh chỉ lựa chọn chính xác 02 ý trong 01 câu hỏi được 0,25 điểm;")
            doc.add_paragraph().add_run("- Thí sinh chỉ lựa chọn chính xác 03 ý trong 01 câu hỏi được 0,5 điểm;")
            doc.add_paragraph().add_run("- Thí sinh lựa chọn chính xác cả 04 ý trong 01 câu hỏi được 1 điểm.")

            # Tạo bảng đáp án đúng sai
            table = doc.add_table(rows=2, cols=len(ds_questions) + 1)

            # Header
            table.cell(0, 0).text = "Câu"
            table.cell(1, 0).text = "Đáp án"

            for i, question in enumerate(ds_questions):
                table.cell(0, i + 1).text = str(i + 1)
                table.cell(0, i + 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Tạo đáp án đúng sai (a) Đúng, b) Sai, c) Đúng, d) Sai)
                dap_an = question.get("dap_an", {})
                answer_text = ""
                for key in ['a', 'b', 'c', 'd']:
                    if key in dap_an:
                        status = "Đúng" if dap_an[key] else "Sai"
                        answer_text += f"{key}) {status}\n"

                table.cell(1, i + 1).text = answer_text.strip()
                table.cell(1, i + 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            logger.error(f"Error creating DS answer section: {e}")

    def _create_dt_answer_section(self, doc: Document, dt_questions: List[Dict[str, Any]]):
        """Tạo phần đáp án trắc nghiệm trả lời ngắn"""
        try:
            doc.add_paragraph()  # Dòng trống

            # Tiêu đề phần
            section3_para = doc.add_paragraph()
            section3_run = section3_para.add_run("PHẦN III. Câu trắc nghiệm trả lời ngắn. ")
            section3_run.bold = True
            section3_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(dt_questions)}")

            note_para = doc.add_paragraph()
            note_para.add_run("Mỗi câu trả lời đúng thí sinh được 0,25 điểm.")

            # Tạo bảng đáp án trả lời ngắn
            table = doc.add_table(rows=2, cols=len(dt_questions) + 1)

            # Header
            table.cell(0, 0).text = "Câu"
            table.cell(1, 0).text = "Đáp án"

            for i, question in enumerate(dt_questions):
                table.cell(0, i + 1).text = str(i + 1)
                table.cell(0, i + 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Lấy đáp án
                dap_an = question.get("dap_an", {})
                correct_answer = dap_an.get("dung", "")
                if not correct_answer:
                    giai_thich = question.get("giai_thich", "")
                    correct_answer = self._extract_correct_answer_from_explanation(giai_thich, dap_an)

                table.cell(1, i + 1).text = correct_answer
                table.cell(1, i + 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            logger.error(f"Error creating DT answer section: {e}")

    def _create_answer_table(self, doc: Document, tn_questions: List[Dict[str, Any]]):
        """Tạo bảng đáp án trắc nghiệm theo format 1 hàng ngang"""
        try:
            if not tn_questions:
                return

            # Tiêu đề đáp án theo chuẩn THPT
            doc.add_paragraph()
            answer_list_title = doc.add_heading("BẢNG ĐÁP ÁN", level=3)
            answer_list_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            answer_list_title_run = answer_list_title.runs[0]
            answer_list_title_run.font.name = "Times New Roman"
            answer_list_title_run.font.size = Pt(14)
            answer_list_title_run.bold = True

            # Tạo format 1 hàng ngang: 1.A  2.B  3.C  4.D  5.A...
            total_questions = len(tn_questions)

            # Tạo bảng với 2 hàng: header và đáp án
            # Số cột = số câu hỏi + 1 (cột đầu cho label)
            table = doc.add_table(rows=2, cols=total_questions + 1)
            table.style = 'Table Grid'

            # Thiết lập độ rộng cột
            for col in table.columns:
                col.width = Inches(0.5)

            # Hàng 1: Header với số câu
            header_row = table.rows[0]
            header_row.cells[0].text = "Câu"
            header_row.cells[0].paragraphs[0].runs[0].font.bold = True
            header_row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for i in range(total_questions):
                header_row.cells[i + 1].text = str(i + 1)
                header_row.cells[i + 1].paragraphs[0].runs[0].font.bold = True
                header_row.cells[i + 1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Hàng 2: Đáp án
            answer_row = table.rows[1]
            answer_row.cells[0].text = "Đáp án"
            answer_row.cells[0].paragraphs[0].runs[0].font.bold = True
            answer_row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Điền đáp án cho từng câu
            for i, question in enumerate(tn_questions):
                dap_an = question.get("dap_an", {})
                correct_answer = dap_an.get("dung", "")

                # Nếu không có đáp án đúng, thử trích xuất từ giải thích
                if not correct_answer:
                    giai_thich = question.get("giai_thich", "")
                    correct_answer = self._extract_correct_answer_from_explanation(giai_thich, dap_an)
                    if not correct_answer:
                        correct_answer = "?"

                answer_row.cells[i + 1].text = correct_answer
                answer_row.cells[i + 1].paragraphs[0].runs[0].font.bold = True
                answer_row.cells[i + 1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Thêm giải thích chi tiết cho từng câu
            doc.add_paragraph()
            detail_title = doc.add_heading("Giải thích chi tiết:", level=3)
            detail_title_run = detail_title.runs[0]
            detail_title_run.font.name = "Times New Roman"
            detail_title_run.font.size = Pt(12)

            for i, question in enumerate(tn_questions, 1):
                # Câu hỏi và đáp án đúng
                answer_para = doc.add_paragraph()
                answer_para.add_run(f"Câu {i}: ").bold = True

                dap_an = question.get("dap_an", {})
                correct_answer = dap_an.get("dung", "")
                answer_para.add_run(f"Đáp án: {correct_answer}")

                # Giải thích
                giai_thich = question.get("giai_thich", "")
                if giai_thich:
                    explain_para = doc.add_paragraph()
                    explain_para.add_run("Giải thích: ").italic = True
                    explain_para.add_run(giai_thich)

                doc.add_paragraph()  # Khoảng trống

        except Exception as e:
            logger.error(f"Error creating answer list: {e}")

    def _create_multi_column_answer_table(self, doc: Document, tn_questions: List[Dict[str, Any]], answers_per_column: int):
        """Tạo bảng đáp án nhiều cột"""
        try:
            total_questions = len(tn_questions)
            num_columns = (total_questions + answers_per_column - 1) // answers_per_column

            # Tạo bảng với số cột phù hợp
            table = doc.add_table(rows=answers_per_column, cols=num_columns)
            table.autofit = False

            # Thiết lập độ rộng cột
            for col in table.columns:
                col.width = Inches(1.5)

            # Điền đáp án vào bảng
            for i, question in enumerate(tn_questions):
                row_idx = i % answers_per_column
                col_idx = i // answers_per_column

                if col_idx < num_columns and row_idx < answers_per_column:
                    cell = table.cell(row_idx, col_idx)
                    cell_para = cell.paragraphs[0]

                    dap_an = question.get("dap_an", {})
                    correct_answer = dap_an.get("dung", "")

                    if not correct_answer:
                        giai_thich = question.get("giai_thich", "")
                        correct_answer = self._extract_correct_answer_from_explanation(giai_thich, dap_an)
                        if not correct_answer:
                            correct_answer = "?"

                    cell_para.add_run(f"{i+1}. ").bold = True
                    cell_para.add_run(correct_answer).bold = True

            # Xóa border của bảng để trông gọn gàng hơn
            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))

        except Exception as e:
            logger.error(f"Error creating multi-column answer table: {e}")

    def _create_detailed_answer(
        self, doc: Document, question: Dict[str, Any], question_num: int
    ):
        """Tạo đáp án chi tiết"""
        try:
            # Số câu
            answer_para = doc.add_paragraph()
            answer_para.add_run(f"Câu {question_num}: ").bold = True

            # Đáp án theo loại
            loai_cau = question.get("loai_cau", "")
            dap_an = question.get("dap_an", {})

            if loai_cau == "DT":
                answer_para.add_run(dap_an.get("dap_an_chinh", ""))
            elif loai_cau == "DS":
                ds_answers = []
                for key in ["a", "b", "c", "d"]:
                    if key in dap_an:
                        status = "Đúng" if dap_an[key] else "Sai"
                        ds_answers.append(f"{key}) {status}")
                answer_para.add_run("; ".join(ds_answers))
            elif loai_cau == "TL":
                y_chinh = dap_an.get("y_chinh", [])
                if y_chinh:
                    answer_para.add_run("\n".join([f"- {y}" for y in y_chinh]))

            # Giải thích
            giai_thich = question.get("giai_thich", "")
            if giai_thich:
                explain_para = doc.add_paragraph()
                explain_para.add_run("Giải thích: ").italic = True
                explain_para.add_run(giai_thich)

            doc.add_paragraph()  # Khoảng trống

        except Exception as e:
            logger.error(f"Error creating detailed answer: {e}")

    def _get_question_type_name(self, loai_cau: str) -> str:
        """Lấy tên đầy đủ của loại câu hỏi"""
        names = {
            "TN": "Trắc nghiệm",
            "DT": "Điền từ",
            "DS": "Đúng/Sai",
            "TL": "Tự luận",
        }
        return names.get(loai_cau, loai_cau)

    def _extract_correct_answer_from_explanation(self, explanation: str, dap_an: dict) -> str:
        """Trích xuất đáp án đúng từ giải thích (copy từ exam_generation_service)"""
        try:
            import re

            if not explanation or not isinstance(dap_an, dict):
                return ""

            explanation_lower = explanation.lower()

            # Tìm các pattern rõ ràng nhất trước
            strong_patterns = [
                r"đáp án ([abcd]) đúng",
                r"đáp án đúng là ([abcd])",
                r"([abcd]) đúng vì",
                r"([abcd]) là đáp án đúng",
                r"([abcd]) đúng",
                r"chọn đáp án ([abcd])",
                r"đáp án:\s*([abcd])",
                r"đáp án\s+([abcd])"
            ]

            for pattern in strong_patterns:
                match = re.search(pattern, explanation_lower)
                if match:
                    answer = match.group(1).upper()
                    if answer in dap_an:
                        return answer

            # Tìm pattern yếu hơn
            weak_patterns = [
                r"đáp án ([abcd])",
                r"chọn ([abcd])",
                r"([abcd])\s*[:\-\.]",
                r"^([abcd])\s"
            ]

            for pattern in weak_patterns:
                match = re.search(pattern, explanation_lower)
                if match:
                    answer = match.group(1).upper()
                    if answer in dap_an:
                        return answer

            # Phân tích nội dung đáp án
            option_scores = {}
            for option, content in dap_an.items():
                if option in ['A', 'B', 'C', 'D'] and isinstance(content, str):
                    content_words = content.lower().split()
                    score = 0
                    for word in content_words:
                        if len(word) > 2 and word in explanation_lower:
                            score += 1
                    option_scores[option] = score

            if option_scores:
                best_option = max(option_scores.keys(), key=lambda x: option_scores[x])
                if option_scores[best_option] > 0:
                    return best_option

            return ""

        except Exception as e:
            logger.error(f"Error extracting correct answer: {e}")
            return ""


# Tạo instance global
exam_docx_service = ExamDocxService()
