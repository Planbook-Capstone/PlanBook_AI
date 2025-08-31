"""
Service tạo file DOCX cho đề thi thông minh theo chuẩn THPT 2025
"""

import logging
import os
import random
import tempfile
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Union

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.smart_exam_models import SmartExamRequest

logger = logging.getLogger(__name__)


class SmartExamDocxService:
    """Service tạo file DOCX cho đề thi thông minh theo chuẩn THPT 2025"""

    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "smart_exams"
        self.temp_dir.mkdir(exist_ok=True)

    async def create_smart_exam_docx(
        self, exam_data: Dict[str, Any], exam_request: Union[SmartExamRequest, Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        Tạo file DOCX cho đề thi thông minh theo chuẩn THPT 2025

        Args:
            exam_data: Dữ liệu đề thi đã được tạo
            exam_request: Request gốc chứa thông tin đề thi

        Returns:
            Dict chứa thông tin file đã tạo
        """
        try:
            # Tạo hoặc lấy mã đề
            exam_code = self._get_or_generate_exam_code(exam_request)

            # Tạo document mới
            doc = Document()

            # Thiết lập style
            self._setup_document_style(doc)

            # Tạo trang bìa với mã đề
            self._create_cover_page(doc, exam_request, exam_data, exam_code)

            # Tạo bảng hóa trị cho môn Hóa học
            subject = self._get_field(exam_request, "subject", "")

            # Kiểm tra cả "hóa" và "hoa" để đảm bảo
            if "hóa" in subject.lower() or "hoa" in subject.lower():
                self._create_chemistry_valence_table(doc, exam_data.get("questions", []))

            # Tạo nội dung đề thi theo 3 phần
            self._create_exam_content_by_parts(doc, exam_data.get("questions", []))

            # Thêm chữ "Hết"
            self._add_exam_ending(doc)

            # Tạo đáp án theo chuẩn THPT 2025
            self._create_thpt_2025_answer_section(doc, exam_data.get("questions", []))

            # Lưu file
            filename = self._generate_filename(exam_request)
            filepath = self.temp_dir / filename
            doc.save(str(filepath))

            logger.info(f"Created smart exam DOCX: {filepath}")
            return {
                "success": True,
                "file_path": str(filepath),
                "filename": filename,
                "file_size": os.path.getsize(filepath)
            }

        except Exception as e:
            logger.error(f"Error creating smart exam DOCX: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    def _get_field(self, exam_request: Union[SmartExamRequest, Dict[str, Any]], field: str, default: Any = None) -> Any:
        """Helper method to get field from either Pydantic model or dict"""
        if isinstance(exam_request, dict):
            return exam_request.get(field, default)
        else:
            return getattr(exam_request, field, default)

    def _normalize_chemistry_format(self, text: str) -> str:
        """
        Chuyển đổi định dạng hóa học từ HTML sang định dạng chuẩn
        VD: <sup>6</sup>Li -> ⁶Li, S<sub>8</sub> -> S₈, Fe<sup>2+</sup> -> Fe²⁺
        """
        if not text:
            return text

        # Chuyển đổi superscript với số và ký hiệu (chỉ số trên)
        sup_pattern = r'<sup>([^<]+)</sup>'
        superscript_map = {
            '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
            '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
            '+': '⁺', '-': '⁻'
        }

        def replace_sup(match):
            content = match.group(1)
            result = ''
            for char in content:
                result += superscript_map.get(char, char)
            return result

        text = re.sub(sup_pattern, replace_sup, text)

        # Chuyển đổi subscript (chỉ số dưới)
        # Chuyển đổi subscript (chỉ số dưới) - bao gồm cả số và chữ cái n, m
        sub_pattern = r'<sub>([\dnm]+)</sub>'
        subscript_map = {
            '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
            '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
            'n': 'ₙ', 'm': 'ₘ' # Thêm n và m
        }

        def replace_sub(match):
            number = match.group(1)
            return ''.join(subscript_map.get(digit, digit) for digit in number)

        text = re.sub(sub_pattern, replace_sub, text)

        # Chuyển đổi các công thức hóa học thô (không có HTML tags)
        # Pattern: chuyển số thành subscript cho tất cả ký hiệu nguyên tố
        # VD: CH3, H2O, C6H12O6, Ca(OH)2, Al2(SO4)3, C2(H2O)2
        # Tất cả số sau dấu ngoặc đóng đều chuyển thành subscript

        # Không cần bảo vệ gì cả - tất cả số đều chuyển thành subscript
        protected_text = text

        # 2. Chuyển đổi subscript cho tất cả ký hiệu nguyên tố
        # Pattern 1: Số ngay sau ký hiệu nguyên tố (VD: H2, O2, Ca2)
        # Pattern 1: Số hoặc n,m ngay sau ký hiệu nguyên tố (VD: H2, C6, Cn)
        # \b để đảm bảo nó là một "từ" riêng biệt, tránh các từ như "Tinh"
        chemistry_pattern = r'\b([A-Z][a-z]?)(\d[\dnm]*|[nm])\b'

        # Pattern 2: Số sau dấu ngoặc đóng (VD: (OH)2, (SO4)3)
        parenthesis_pattern = r'\)([\dnm]+)'

        def replace_chemistry(match):
            element = match.group(1)
            number = match.group(2)
            # Chuyển số thành subscript
            subscript_number = ''.join(subscript_map.get(digit, digit) for digit in number)
            return element + subscript_number

        def replace_parenthesis(match):
            number = match.group(1)
            # Chuyển số thành subscript
            subscript_number = ''.join(subscript_map.get(digit, digit) for digit in number)
            return ')' + subscript_number

        # Áp dụng cả hai pattern
        protected_text = re.sub(chemistry_pattern, replace_chemistry, protected_text)
        text = re.sub(parenthesis_pattern, replace_parenthesis, protected_text)

        return text

    def _extract_numeric_answer(self, answer_text: str) -> str:
        """
        Trích xuất đáp án số thuần túy từ text cho phần III theo quy tắc mới

        🎯 Quy tắc format đáp án phần 3:
        1. Chỉ được tô đúng 4 ký tự (bao gồm chữ số 0-9, dấu trừ -, dấu phẩy ,)
        2. Làm tròn/cắt bớt phần thập phân sao cho vừa đủ 4 ký tự
        3. Không được tô dư, không ghi dấu , ở cuối
        4. Chuyển đổi dấu thập phân: . → ,

        📘 Ví dụ chuẩn:
        - 12.34 → 12,3 (cắt còn 4 ký tự)
        - -1.56 → -1,5 (đủ 4 ký tự)
        - 0.123 → 0,12 (cắt, đổi . → ,)
        - 123.45 → 123 (ưu tiên phần nguyên)
        - 3.5 → 3,5 (dưới 4 ký tự, đổi dấu)
        - -12.34 → -12 (cắt phần thập phân)
        """
        if not answer_text:
            return "0"

        # Tìm số đầu tiên trong chuỗi
        number_pattern = r'-?(?:\d+[.,]?\d*|[.,]\d+)'
        match = re.search(number_pattern, str(answer_text))

        if match:
            number = match.group()
            try:
                # Chuyển dấu phẩy thành dấu chấm để parse
                number_for_parse = number.replace(',', '.')
                float_num = float(number_for_parse)

                # Xử lý theo quy tắc mới
                if float_num == int(float_num):
                    # Số nguyên - kiểm tra độ dài
                    result = str(int(float_num))
                    if len(result) > 4:
                        # Cắt bớt nếu quá dài
                        result = result[:4]
                else:
                    # Số thập phân - áp dụng quy tắc format
                    if float_num < 0:
                        # Số âm
                        if abs(float_num) >= 100:
                            # VD: -123.45 → -123 (ưu tiên phần nguyên)
                            result = str(int(float_num))[:4]
                        elif abs(float_num) >= 10:
                            # VD: -12.34 → -12 (cắt phần thập phân để vừa 4 ký tự)
                            result = str(int(float_num))
                        else:
                            # VD: -1.56 → -1,5 (đủ 4 ký tự)
                            # VD: -0.5 → -0,5 (số âm nhỏ)
                            if abs(float_num) < 1:
                                # Trường hợp đặc biệt: -0.5 → -0,5
                                decimal_str = f"{abs(float_num):.10f}"[2:]  # Bỏ "0."
                                available_chars = 4 - 2 - 1  # -0, = 3 ký tự đã dùng
                                if available_chars > 0:
                                    decimal_truncated = decimal_str[:available_chars]
                                    result = f"-0,{decimal_truncated}"
                                else:
                                    result = "-0"
                            else:
                                integer_part = int(float_num)
                                decimal_part = abs(float_num) - abs(integer_part)

                                # Tính số chữ số thập phân có thể có
                                available_chars = 4 - len(str(integer_part)) - 1  # -1 cho dấu phẩy
                                if available_chars > 0:
                                    # Lấy phần thập phân và cắt theo số ký tự có thể
                                    decimal_str = f"{decimal_part:.10f}"[2:]  # Bỏ "0."
                                    decimal_truncated = decimal_str[:available_chars]
                                    result = f"{integer_part},{decimal_truncated}"
                                else:
                                    result = str(integer_part)
                    else:
                        # Số dương
                        if float_num >= 1000:
                            # VD: 1234.56 → 1234 (ưu tiên phần nguyên)
                            result = str(int(float_num))[:4]
                        elif float_num >= 100:
                            # VD: 123.45 → 123 (ưu tiên phần nguyên)
                            result = str(int(float_num))
                        elif float_num >= 10:
                            # VD: 12.34 → 12,3 (cắt còn 4 ký tự)
                            integer_part = int(float_num)
                            decimal_part = float_num - integer_part

                            # Có thể có 1 chữ số thập phân (XX,Y = 4 ký tự)
                            decimal_str = f"{decimal_part:.10f}"[2:]  # Bỏ "0."
                            result = f"{integer_part},{decimal_str[0]}"
                        elif float_num >= 1:
                            # VD: 3.5 → 3,5 (dưới 4 ký tự)
                            integer_part = int(float_num)
                            decimal_part = float_num - integer_part

                            # Có thể có 2 chữ số thập phân (X,YZ = 4 ký tự)
                            available_chars = 4 - len(str(integer_part)) - 1  # -1 cho dấu phẩy
                            decimal_str = f"{decimal_part:.10f}"[2:]  # Bỏ "0."
                            decimal_truncated = decimal_str[:available_chars]
                            result = f"{integer_part},{decimal_truncated}"
                        else:
                            # VD: 0.123 → 0,12 (cắt, đổi . → ,)
                            # VD: 0.0025 → 0,00 (số rất nhỏ)
                            # VD: 0.9 → 0,9 (không thêm số 0 thừa)
                            decimal_str = f"{float_num:.10f}"[2:]  # Bỏ "0."
                            # Có thể có 2 chữ số thập phân (0,YZ = 4 ký tự)
                            if len(decimal_str) >= 2:
                                result = f"0,{decimal_str[:2]}"
                            else:
                                result = f"0,{decimal_str}"  # Không thêm số 0 thừa

                # Loại bỏ dấu phẩy ở cuối và số 0 thừa (nhưng giữ lại số 0 có ý nghĩa)
                if result.endswith(','):
                    result = result[:-1]
                elif ',' in result:
                    # Loại bỏ số 0 thừa ở cuối, nhưng giữ lại ít nhất 1 chữ số sau dấu phẩy
                    # Trường hợp đặc biệt: 0,00 (số rất nhỏ) thì giữ nguyên
                    if result.startswith('0,00'):
                        pass  # Giữ nguyên 0,00
                    else:
                        # Loại bỏ số 0 thừa: 0,90 → 0,9, 1,50 → 1,5
                        result = result.rstrip('0')
                        if result.endswith(','):
                            result = result[:-1]

                # Đảm bảo không vượt quá 4 ký tự
                if len(result) > 4:
                    result = result[:4]
                    if result.endswith(','):
                        result = result[:-1]

                return result

            except ValueError:
                return "0"

        return "0"

    def _get_or_generate_exam_code(self, exam_request: Union[SmartExamRequest, Dict[str, Any]]) -> str:
        """Lấy mã đề từ request hoặc tạo random nếu không có"""
        try:
            exam_code = self._get_field(exam_request, "examCode", None)
            if exam_code:
                return exam_code

            # Tạo mã đề random 4 số
            return f"{random.randint(1000, 9999)}"

        except Exception as e:
            logger.error(f"Error getting or generating exam code: {e}")
            return f"{random.randint(1000, 9999)}"

    def _setup_document_style(self, doc: Document):
        """Thiết lập style cho document"""
        try:
            # Thiết lập margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            # Thiết lập font mặc định
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)

        except Exception as e:
            logger.error(f"Error setting up document style: {e}")

    def _create_cover_page(self, doc: Document, exam_request: Union[SmartExamRequest, Dict[str, Any]], exam_data: Dict[str, Any], exam_code: str):
        """Tạo trang bìa theo chuẩn THPT 2025 với mã đề"""
        try:
            # Header với 2 cột: thông tin trường, thông tin đề thi
            header_table = doc.add_table(rows=1, cols=2)
            header_table.autofit = False

            # Thiết lập độ rộng cột
            header_table.columns[0].width = Inches(3.0)  # Cột trái - thông tin trường
            header_table.columns[1].width = Inches(4.0)  # Cột phải - thông tin đề thi

            # Cột trái - Logo và thông tin bộ
            left_cell = header_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            left_para.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO").bold = True
            left_para.add_run("\n")
            left_para.add_run(self._get_field(exam_request, "school", "TRƯỜNG THPT ABC")).bold = True

            # Cột phải - Thông tin đề thi
            right_cell = header_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Lấy grade từ request để hiển thị đúng lớp
            grade = self._get_field(exam_request, "grade", 10)
            right_para.add_run(f"ĐỀ KIỂM TRA LỚP {grade}").bold = True
            right_para.add_run(f"\nMôn: {self._get_field(exam_request, 'subject', 'HÓA HỌC').upper()}")
            right_para.add_run(f"\nThời gian làm bài: {self._get_field(exam_request, 'duration', 50)} phút, không kể thời gian phát đề")

            # Loại bỏ border cho cả 2 ô
            left_cell._element.get_or_add_tcPr().append(
                self._create_no_border_element()
            )
            right_cell._element.get_or_add_tcPr().append(
                self._create_no_border_element()
            )

            # Khoảng trắng sau header
            doc.add_paragraph()

            # Thông tin thí sinh với mã đề cùng hàng
            info_table = doc.add_table(rows=2, cols=2)
            info_table.autofit = False

            # Thiết lập độ rộng cột cho bảng thông tin
            info_table.columns[0].width = Inches(4.5)  # Cột trái - thông tin thí sinh
            info_table.columns[1].width = Inches(2.5)  # Cột phải - mã đề

            # Hàng 1: Họ tên thí sinh và mã đề
            name_cell = info_table.cell(0, 0)
            name_para = name_cell.paragraphs[0]
            name_para.add_run("Họ, tên thí sinh: ").bold = True
            name_para.add_run("." * 50)

            # Ô mã đề cùng hàng với họ tên
            code_cell = info_table.cell(0, 1)
            code_para = code_cell.paragraphs[0]
            code_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Nội dung mã đề
            ma_de_run = code_para.add_run("Mã đề: ")
            ma_de_run.bold = True
            ma_de_run.font.size = Pt(11)
            code_run = code_para.add_run(exam_code)
            code_run.bold = True
            code_run.font.size = Pt(12)

            # Thiết lập border cho ô mã đề
            self._set_cell_border_enhanced(code_cell)

            # Hàng 2: Số báo danh (ô thứ 2 để trống)
            sbd_cell = info_table.cell(1, 0)
            sbd_para = sbd_cell.paragraphs[0]
            sbd_para.add_run("Số báo danh: ").bold = True
            sbd_para.add_run("." * 55)

            # Ô trống bên phải số báo danh
            empty_cell = info_table.cell(1, 1)

            # Loại bỏ border cho các ô không phải mã đề
            name_cell._element.get_or_add_tcPr().append(self._create_no_border_element())
            sbd_cell._element.get_or_add_tcPr().append(self._create_no_border_element())
            empty_cell._element.get_or_add_tcPr().append(self._create_no_border_element())

            # Thống kê đề thi
            doc.add_paragraph()
            statistics = exam_data.get("statistics", {})
            if hasattr(statistics, 'total_questions'):
                total_questions = statistics.total_questions
            elif isinstance(statistics, dict):
                total_questions = statistics.get("total_questions", 0)
            else:
                total_questions = 0
            
            stats_para = doc.add_paragraph()
            stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            stats_para.add_run(f"(Đề thi gồm {total_questions} câu)")


        except Exception as e:
            logger.error(f"Error creating cover page: {e}")

   
    def _create_chemistry_valence_table(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo bảng nguyên tử khối cho môn Hóa học dựa trên nội dung đề thi"""
        try:
            # Phân tích nội dung đề thi để tìm các nguyên tố
            used_elements = self._extract_chemical_elements_from_questions(questions)

            if not used_elements:
                # Nếu không tìm thấy nguyên tố nào, sử dụng các nguyên tố phổ biến cho hóa học THPT
                used_elements = ["H", "C", "N", "O", "Na", "Mg", "Al", "Si", "P", "S", "Cl", "K", "Ca", "Fe", "Cu", "Zn", "Br", "I", "Ag", "Ba"]

            # Tiêu đề
            valence_title = doc.add_paragraph()
            valence_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            valence_run = valence_title.add_run("BẢNG NGUYÊN TỬ KHỐI CỦA CÁC NGUYÊN TỐ HÓA HỌC")
            valence_run.bold = True

            # Tạo bảng nguyên tử khối chỉ cho các nguyên tố được sử dụng
            atomic_masses_text = self._get_atomic_masses_for_elements(used_elements)

            valence_para = doc.add_paragraph()
            valence_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            valence_para.add_run(atomic_masses_text)

            # Thêm lưu ý
            # note_para = doc.add_paragraph()
            # note_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # if len(used_elements) <= 20:  # Nếu ít nguyên tố thì ghi chú là từ đề thi
            #     # note_run = note_para.add_run("(Chỉ ghi các nguyên tố có trong đề thi)")
            #     print("DEBUG: Not enough elements to create custom valence table")
            # else:  # Nếu nhiều nguyên tố thì ghi chú là nguyên tố phổ biến
            #     note_run = note_para.add_run("(Các nguyên tố hóa học phổ biến)")
            # note_run.italic = True
            # note_run.font.size = Pt(10)

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating chemistry valence table: {e}")

    def _extract_chemical_elements_from_questions(self, questions: List[Dict[str, Any]]) -> List[str]:
        """Trích xuất các nguyên tố hóa học từ nội dung câu hỏi"""
        import re

        # Danh sách các nguyên tố hóa học phổ biến
        all_elements = {
            "H": 1, "He": 4, "Li": 7, "Be": 9, "B": 11, "C": 12, "N": 14, "O": 16,
            "F": 19, "Ne": 20, "Na": 23, "Mg": 24, "Al": 27, "Si": 28, "P": 31, "S": 32,
            "Cl": 35.5, "Ar": 40, "K": 39, "Ca": 40, "Sc": 45, "Ti": 48, "V": 51, "Cr": 52,
            "Mn": 55, "Fe": 56, "Co": 59, "Ni": 59, "Cu": 64, "Zn": 65, "Ga": 70, "Ge": 73,
            "As": 75, "Se": 79, "Br": 80, "Kr": 84, "Rb": 85, "Sr": 88, "Y": 89, "Zr": 91,
            "Nb": 93, "Mo": 96, "Tc": 98, "Ru": 101, "Rh": 103, "Pd": 106, "Ag": 108, "Cd": 112,
            "In": 115, "Sn": 119, "Sb": 122, "Te": 128, "I": 127, "Xe": 131, "Cs": 133, "Ba": 137,
            "La": 139, "Ce": 140, "Pr": 141, "Nd": 144, "Pm": 145, "Sm": 150, "Eu": 152, "Gd": 157,
            "Tb": 159, "Dy": 163, "Ho": 165, "Er": 167, "Tm": 169, "Yb": 173, "Lu": 175, "Hf": 178,
            "Ta": 181, "W": 184, "Re": 186, "Os": 190, "Ir": 192, "Pt": 195, "Au": 197, "Hg": 201,
            "Tl": 204, "Pb": 207, "Bi": 209, "Po": 209, "At": 210, "Rn": 222, "Fr": 223, "Ra": 226,
            "Ac": 227, "Th": 232, "Pa": 231, "U": 238, "Np": 237, "Pu": 244, "Am": 243, "Cm": 247,
            "Bk": 247, "Cf": 251, "Es": 252, "Fm": 257, "Md": 258, "No": 259, "Lr": 262
        }

        found_elements = set()

        # Tìm kiếm trong tất cả nội dung câu hỏi
        for i, question in enumerate(questions):
            question_text = str(question.get("question", ""))
            answer_text = str(question.get("answer", ""))
            explanation_text = str(question.get("explanation", ""))

            # Kết hợp tất cả text
            full_text = f"{question_text} {answer_text} {explanation_text}"

            # Tìm các ký hiệu nguyên tố với pattern chính xác hơn
            # Tìm các pattern như: H2O, NaCl, CaCO3, Fe2O3, etc.
            element_patterns = [
                r'\b([A-Z][a-z]?)(?:\d+)?',  # H, H2, Na, Cl, etc.
                r'([A-Z][a-z]?)(?=\d)',      # H trong H2O, Na trong NaCl
                r'([A-Z][a-z]?)(?=[A-Z])',   # Na trong NaCl, Ca trong CaCO3
                r'\b([A-Z][a-z]?)\b'         # Standalone elements
            ]

            question_elements = set()
            for pattern in element_patterns:
                matches = re.findall(pattern, full_text)
                for match in matches:
                    if match in all_elements:
                        found_elements.add(match)
                        question_elements.add(match)
        return sorted(list(found_elements))

    def _get_atomic_masses_for_elements(self, elements: List[str]) -> str:
        """Tạo chuỗi nguyên tử khối cho các nguyên tố được chỉ định"""
        atomic_masses = {
            "H": 1, "He": 4, "Li": 7, "Be": 9, "B": 11, "C": 12, "N": 14, "O": 16,
            "F": 19, "Ne": 20, "Na": 23, "Mg": 24, "Al": 27, "Si": 28, "P": 31, "S": 32,
            "Cl": 35.5, "Ar": 40, "K": 39, "Ca": 40, "Sc": 45, "Ti": 48, "V": 51, "Cr": 52,
            "Mn": 55, "Fe": 56, "Co": 59, "Ni": 59, "Cu": 64, "Zn": 65, "Ga": 70, "Ge": 73,
            "As": 75, "Se": 79, "Br": 80, "Kr": 84, "Rb": 85, "Sr": 88, "Y": 89, "Zr": 91,
            "Nb": 93, "Mo": 96, "Tc": 98, "Ru": 101, "Rh": 103, "Pd": 106, "Ag": 108, "Cd": 112,
            "In": 115, "Sn": 119, "Sb": 122, "Te": 128, "I": 127, "Xe": 131, "Cs": 133, "Ba": 137,
            "La": 139, "Ce": 140, "Pr": 141, "Nd": 144, "Pm": 145, "Sm": 150, "Eu": 152, "Gd": 157,
            "Tb": 159, "Dy": 163, "Ho": 165, "Er": 167, "Tm": 169, "Yb": 173, "Lu": 175, "Hf": 178,
            "Ta": 181, "W": 184, "Re": 186, "Os": 190, "Ir": 192, "Pt": 195, "Au": 197, "Hg": 201,
            "Tl": 204, "Pb": 207, "Bi": 209, "Po": 209, "At": 210, "Rn": 222, "Fr": 223, "Ra": 226,
            "Ac": 227, "Th": 232, "Pa": 231, "U": 238, "Np": 237, "Pu": 244, "Am": 243, "Cm": 247,
            "Bk": 247, "Cf": 251, "Es": 252, "Fm": 257, "Md": 258, "No": 259, "Lr": 262
        }

        mass_strings = []
        for element in elements:
            if element in atomic_masses:
                mass = atomic_masses[element]
                if mass == int(mass):
                    mass_strings.append(f"{element} = {int(mass)}")
                else:
                    mass_strings.append(f"{element} = {mass}")

        return "; ".join(mass_strings)

    def _create_exam_content_by_parts(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo nội dung đề thi theo 3 phần chuẩn THPT 2025"""
        try:
            print(f"DEBUG DOCX: Total questions received: {len(questions)}")
            for i, q in enumerate(questions[:3]):  # Show first 3 questions
                print(f"DEBUG DOCX: Question {i+1} keys: {list(q.keys())}")
                print(f"DEBUG DOCX: Question {i+1} part: {q.get('part', 'NO_PART')}")
                print(f"DEBUG DOCX: Question {i+1} answer structure: {q.get('answer', 'NO_ANSWER')}")
                print(f"DEBUG DOCX: Question {i+1} dap_an structure: {q.get('dap_an', 'NO_DAP_AN')}")
                if q.get('answer'):
                    print(f"DEBUG DOCX: Question {i+1} answer keys: {list(q.get('answer', {}).keys())}")
                if q.get('dap_an'):
                    print(f"DEBUG DOCX: Question {i+1} dap_an keys: {list(q.get('dap_an', {}).keys())}")

            # Phân loại câu hỏi theo phần
            part_1_questions = [q for q in questions if q.get("part") == 1]
            part_2_questions = [q for q in questions if q.get("part") == 2]
            part_3_questions = [q for q in questions if q.get("part") == 3]

            print(f"DEBUG DOCX: Part 1 questions: {len(part_1_questions)}")
            print(f"DEBUG DOCX: Part 2 questions: {len(part_2_questions)}")
            print(f"DEBUG DOCX: Part 3 questions: {len(part_3_questions)}")

            # PHẦN I: Câu trắc nghiệm nhiều phương án lựa chọn
            if part_1_questions:
                self._create_part_1_section(doc, part_1_questions)

            # PHẦN II: Câu trắc nghiệm đúng sai
            if part_2_questions:
                self._create_part_2_section(doc, part_2_questions)

            # PHẦN III: Câu trắc nghiệm trả lời ngắn
            if part_3_questions:
                self._create_part_3_section(doc, part_3_questions)

        except Exception as e:
            logger.error(f"Error creating exam content by parts: {e}")

    def _create_part_1_section(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo PHẦN I: Câu trắc nghiệm nhiều phương án lựa chọn"""
        try:
            # Tiêu đề phần
            part_title = doc.add_paragraph()
            part_title_run = part_title.add_run("PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. ")
            part_title_run.bold = True
            part_title.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}.")

            note_para = doc.add_paragraph()
            note_para.add_run("(Mỗi câu trả lời đúng thí sinh được 0,25 điểm)")

            # Tạo câu hỏi
            for i, question in enumerate(questions, 1):
                self._create_multiple_choice_question(doc, question, i)

        except Exception as e:
            logger.error(f"Error creating part 1 section: {e}")

    def _create_part_2_section(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo PHẦN II: Câu trắc nghiệm đúng sai"""
        try:
            # Tiêu đề phần
            part_title = doc.add_paragraph()
            part_title_run = part_title.add_run("PHẦN II. Câu trắc nghiệm đúng sai. ")
            part_title_run.bold = True
            part_title.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.")

            # Hướng dẫn chấm điểm
   

            # Tạo câu hỏi
            for i, question in enumerate(questions, 1):
                self._create_true_false_question_with_statements(doc, question, i)

        except Exception as e:
            logger.error(f"Error creating part 2 section: {e}")

    def _create_part_3_section(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo PHẦN III: Câu trắc nghiệm trả lời ngắn"""
        try:
            # Tiêu đề phần
            part_title = doc.add_paragraph()
            part_title_run = part_title.add_run("PHẦN III. Câu trắc nghiệm trả lời ngắn. ")
            part_title_run.bold = True
            part_title.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}")

            note_para = doc.add_paragraph()
            # Thêm lưu ý về định dạng đáp án
            format_note = doc.add_paragraph()
            format_note_run = format_note.add_run("Lưu ý: ")
            format_note_run.bold = True
            format_note.add_run("Đáp án phần III chỉ ghi số (không ghi đơn vị, không ghi chữ). ")
            format_note.add_run("Sử dụng dấu phẩy (,), tối đa 4 ký tự. ")
            format_note.add_run("Đáp án chỉ lấy số nguyên không tính phần lẻ, học sinh tự làm tròn. ")
            format_note.add_run("Ví dụ: 12,3; -1,5; 0,12; 123")

            doc.add_paragraph()

            # Tạo câu hỏi
            for i, question in enumerate(questions, 1):
                self._create_short_answer_question(doc, question, i)

        except Exception as e:
            logger.error(f"Error creating part 3 section: {e}")

    def _create_multiple_choice_question(self, doc: Document, question: Dict[str, Any], question_num: int):
        """Tạo câu hỏi trắc nghiệm nhiều phương án với layout thông minh"""
        try:
            # Câu hỏi
            q_para = doc.add_paragraph()
            q_para.add_run(f"Câu {question_num}. ").bold = True
            question_text = question.get("question", question.get("cau_hoi", ""))
            q_para.add_run(self._normalize_chemistry_format(question_text))

            # Lấy các phương án
            dap_an = question.get("answer", question.get("dap_an", {}))
            options = {}
            for option in ["A", "B", "C", "D"]:
                if option in dap_an:
                    options[option] = self._normalize_chemistry_format(str(dap_an[option]))

            # Quyết định layout dựa trên độ dài đáp án và tạo hiển thị
            self._create_options_with_smart_layout(doc, options)

        except Exception as e:
            logger.error(f"Error creating multiple choice question: {e}")

    def _create_options_with_smart_layout(self, doc: Document, options: Dict[str, str]):
        """Tạo các lựa chọn với layout thông minh dựa trên độ dài nội dung"""
        try:
            if not options:
                return

            # Quyết định layout dựa trên độ dài đáp án
            layout = self._determine_options_layout(options)
            
            if layout == "single_row":
                self._create_options_single_row(doc, options)
            elif layout == "double_row":
                self._create_options_double_row(doc, options)
            else:  # four_rows
                self._create_options_four_rows(doc, options)

        except Exception as e:
            logger.error(f"Error creating options with smart layout: {e}")

    def _determine_options_layout(self, options: Dict[str, str]) -> str:
        """
        Quyết định layout hiển thị các lựa chọn dựa trên độ dài nội dung
        
        Logic:
        - Nếu tất cả đáp án ngắn (≤ 25 ký tự): 1 hàng
        - Nếu đáp án vừa phải (26-60 ký tự): 2 hàng (mỗi hàng 2 đáp án)  
        - Nếu có đáp án dài (> 60 ký tự): 4 hàng (mỗi đáp án 1 hàng)
        """
        if not options:
            return "four_rows"

        max_length = 0
        total_length = 0
        
        for option_text in options.values():
            length = len(option_text.strip())
            max_length = max(max_length, length)
            total_length += length
        
        avg_length = total_length / len(options)
        
        # Logic quyết định layout
        if max_length <= 25 and avg_length <= 20:
            # Tất cả đáp án ngắn -> 1 hàng
            return "single_row"
        elif max_length <= 60 and avg_length <= 45:
            # Đáp án vừa phải -> 2 hàng
            return "double_row"
        else:
            # Có đáp án dài -> 4 hàng (giữ nguyên format cũ)
            return "four_rows"

    def _create_options_single_row(self, doc: Document, options: Dict[str, str]):
        """Tạo các lựa chọn trên 1 hàng với căn lề trái và giãn đều bằng bảng"""
        try:
            # Đếm số lựa chọn có sẵn
            available_options = [option for option in ["A", "B", "C", "D"] if option in options]
            if not available_options:
                return

            # Tạo bảng với số cột bằng số lựa chọn
            table = doc.add_table(rows=1, cols=len(available_options))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Loại bỏ border của bảng để trông như text bình thường
            for row in table.rows:
                for cell in row.cells:
                    # Xóa border
                    cell._element.get_or_add_tcPr().append(
                        self._create_no_border_element()
                    )
                    # Căn lề trái nội dung trong cell
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Điền nội dung vào các cell
            for i, option in enumerate(available_options):
                cell = table.cell(0, i)
                cell.text = f"{option}. {options[option]}"
                # Đảm bảo căn lề trái
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        except Exception as e:
            logger.error(f"Error creating single row options: {e}")

    def _create_no_border_element(self):
        """Tạo element XML để loại bỏ border của table cell"""
        try:
            from docx.oxml import parse_xml
            no_border_xml = """
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="nil"/>
                <w:left w:val="nil"/>
                <w:bottom w:val="nil"/>
                <w:right w:val="nil"/>
            </w:tcBorders>
            """
            return parse_xml(no_border_xml)
        except Exception as e:
            logger.error(f"Error creating no border element: {e}")
            return None

    def _create_exam_code_border(self):
        """Tạo element XML để tạo border đậm cho ô mã đề"""
        try:
            from docx.oxml import parse_xml
            border_xml = """
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>
            </w:tcBorders>
            """
            return parse_xml(border_xml)
        except Exception as e:
            logger.error(f"Error creating exam code border element: {e}")
            return None

    def _create_cell_padding(self):
        """Tạo element XML để thêm padding cho cell"""
        try:
            from docx.oxml import parse_xml
            padding_xml = """
            <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:w="100" w:type="dxa"/>
                <w:left w:w="100" w:type="dxa"/>
                <w:bottom w:w="100" w:type="dxa"/>
                <w:right w:w="100" w:type="dxa"/>
            </w:tcMar>
            """
            return parse_xml(padding_xml)
        except Exception as e:
            logger.error(f"Error creating cell padding element: {e}")
            return None

    def _set_cell_border(self, cell):
        """Thiết lập border rõ ràng cho cell"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Tạo tcPr element nếu chưa có
            tcPr = cell._element.get_or_add_tcPr()

            # Tạo tcBorders element
            tcBorders = OxmlElement('w:tcBorders')

            # Thiết lập border cho 4 phía
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')  # Độ dày border
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Màu đen
                tcBorders.append(border)

            # Thêm tcBorders vào tcPr
            tcPr.append(tcBorders)

            # Thêm padding
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '100')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)

            tcPr.append(tcMar)

        except Exception as e:
            logger.error(f"Error setting cell border: {e}")

    def _set_cell_border_enhanced(self, cell):
        """Thiết lập border rất mỏng và padding tối thiểu cho ô mã đề"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Tạo tcPr element nếu chưa có
            tcPr = cell._element.get_or_add_tcPr()

            # Xóa border cũ nếu có
            existing_borders = tcPr.find(qn('w:tcBorders'))
            if existing_borders is not None:
                tcPr.remove(existing_borders)

            # Tạo tcBorders element mới
            tcBorders = OxmlElement('w:tcBorders')

            # Thiết lập border cực mỏng cho 4 phía
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '2')  # Border cực mỏng
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Màu đen
                tcBorders.append(border)

            # Thêm tcBorders vào tcPr
            tcPr.append(tcBorders)

            # Thêm padding tối thiểu
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '20')  # Padding tối thiểu (giảm từ 40 xuống 20)
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)

            tcPr.append(tcMar)

            # Không thêm background color (loại bỏ màu nền)

        except Exception as e:
            logger.error(f"Error setting enhanced cell border: {e}")

    def _create_options_double_row(self, doc: Document, options: Dict[str, str]):
        """Tạo các lựa chọn trên 2 hàng với căn lề trái và giãn đều bằng bảng"""
        try:
            # Hàng 1: A và B
            row1_options = []
            for option in ["A", "B"]:
                if option in options:
                    row1_options.append((option, options[option]))

            if row1_options:
                # Tạo bảng cho hàng 1
                table1 = doc.add_table(rows=1, cols=len(row1_options))
                table1.alignment = WD_TABLE_ALIGNMENT.LEFT

                # Loại bỏ border và căn lề trái
                for i, (option, text) in enumerate(row1_options):
                    cell = table1.cell(0, i)
                    cell._element.get_or_add_tcPr().append(
                        self._create_no_border_element()
                    )
                    cell.text = f"{option}. {text}"
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Hàng 2: C và D
            row2_options = []
            for option in ["C", "D"]:
                if option in options:
                    row2_options.append((option, options[option]))

            if row2_options:
                # Tạo bảng cho hàng 2
                table2 = doc.add_table(rows=1, cols=len(row2_options))
                table2.alignment = WD_TABLE_ALIGNMENT.LEFT

                # Loại bỏ border và căn lề trái
                for i, (option, text) in enumerate(row2_options):
                    cell = table2.cell(0, i)
                    cell._element.get_or_add_tcPr().append(
                        self._create_no_border_element()
                    )
                    cell.text = f"{option}. {text}"
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        except Exception as e:
            logger.error(f"Error creating double row options: {e}")

    def _create_options_four_rows(self, doc: Document, options: Dict[str, str]):
        """Tạo các lựa chọn trên 4 hàng (mỗi đáp án 1 hàng - format gốc cho đáp án dài)"""
        try:
            for option in ["A", "B", "C", "D"]:
                if option in options:
                    option_para = doc.add_paragraph()
                    option_para.add_run(f"{option}. {options[option]}")
                    # Thêm căn lề trái nhẹ để đồng nhất với các đáp án khác
                    option_para.paragraph_format.left_indent = Inches(0.2)

        except Exception as e:
            logger.error(f"Error creating four rows options: {e}")

    def _create_true_false_question(self, doc: Document, question: Dict[str, Any], question_num: int):
        """Tạo câu hỏi đúng sai"""
        try:
            # Câu hỏi chính
            q_para = doc.add_paragraph()
            q_para.add_run(f"Câu {question_num}. ").bold = True
            # Sử dụng field "question" thay vì "cau_hoi"
            q_para.add_run(question.get("question", question.get("cau_hoi", "")))

            # Các ý a, b, c, d
            # Sử dụng field "answer" thay vì "dap_an"
            dap_an = question.get("answer", question.get("dap_an", {}))
            for option in ["a", "b", "c", "d"]:
                if option in dap_an:
                    option_para = doc.add_paragraph()
                    option_para.add_run(f"{option}) {dap_an[option]}")

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating true false question: {e}")

    def _create_true_false_question_with_statements(self, doc: Document, question: Dict[str, Any], question_num: int):
        """Tạo câu hỏi đúng sai với các statement a), b), c), d) theo mẫu THPT"""
        try:
            # Câu hỏi chính
            q_para = doc.add_paragraph()
            q_para.add_run(f"Câu {question_num}. ").bold = True
            # Sử dụng field "question" thay vì "cau_hoi" và chuẩn hóa định dạng hóa học
            main_question = question.get("question", question.get("cau_hoi", ""))
            q_para.add_run(self._normalize_chemistry_format(main_question))

            # Các statement a), b), c), d) - lấy từ explanation hoặc tạo từ answer
            answer_data = question.get("answer", question.get("dap_an", {}))
            explanation = question.get("explanation", "")

            # Tạo các statement từ explanation hoặc answer
            statements = self._extract_statements_from_question(answer_data, explanation)

            for option in ["a", "b", "c", "d"]:
                if option in statements:
                    option_para = doc.add_paragraph()
                    statement_text = self._normalize_chemistry_format(statements[option])
                    option_para.add_run(f"{option}) {statement_text}")

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating true false question with statements: {e}")

    def _extract_statements_from_question(self, answer_data: Dict[str, Any], explanation: str) -> Dict[str, str]:
        """Trích xuất các statement từ dữ liệu câu hỏi"""
        try:
            statements = {}

            # Nếu answer_data đã có các statement a, b, c, d
            for option in ["a", "b", "c", "d"]:
                if option in answer_data:
                    option_data = answer_data[option]

                    # Xử lý cấu trúc mới với content và evaluation
                    if isinstance(option_data, dict) and "content" in option_data:
                        statements[option] = option_data["content"]
                    elif isinstance(option_data, str):
                        # Nếu là đáp án Đúng/Sai, tạo statement mặc định
                        if option_data in ["Đúng", "Sai"]:
                            statements[option] = f"Statement {option.upper()}"
                        else:
                            # Nếu là statement thực tế (cấu trúc cũ)
                            statements[option] = option_data

            # Nếu không có statement, tạo mặc định
            if not statements:
                statements = {
                    "a": "Statement A",
                    "b": "Statement B",
                    "c": "Statement C",
                    "d": "Statement D"
                }

            return statements

        except Exception as e:
            logger.error(f"Error extracting statements: {e}")
            return {"a": "Statement A", "b": "Statement B", "c": "Statement C", "d": "Statement D"}

    def _create_short_answer_question(self, doc: Document, question: Dict[str, Any], question_num: int):
        """Tạo câu hỏi trả lời ngắn"""
        try:
            # Câu hỏi
            q_para = doc.add_paragraph()
            q_para.add_run(f"Câu {question_num}. ").bold = True
            # Sử dụng field "question" thay vì "cau_hoi" và chuẩn hóa định dạng hóa học
            question_text = question.get("question", question.get("cau_hoi", ""))
            q_para.add_run(self._normalize_chemistry_format(question_text))

            doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating short answer question: {e}")

    def _add_exam_ending(self, doc: Document):
        """Thêm chữ 'Hết' vào cuối đề thi"""
        try:
            doc.add_paragraph()
            ending_para = doc.add_paragraph()
            ending_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            ending_run = ending_para.add_run("----- HẾT -----")
            ending_run.bold = True
            ending_run.font.size = Pt(14)
            doc.add_page_break()
        except Exception as e:
            logger.error(f"Error adding exam ending: {e}")

    def _create_thpt_2025_answer_section(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo phần đáp án theo chuẩn THPT 2025"""
        try:
            doc.add_paragraph()
            doc.add_paragraph()

            # Tiêu đề đáp án
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run("ĐÁP ÁN")
            title_run.bold = True
            title_run.font.size = Pt(14)

            doc.add_paragraph()

            # Phân loại câu hỏi theo phần
            part_1_questions = [q for q in questions if q.get("part") == 1]
            part_2_questions = [q for q in questions if q.get("part") == 2]
            part_3_questions = [q for q in questions if q.get("part") == 3]

            # Tạo đáp án cho từng phần
            if part_1_questions:
                self._create_part_1_answer_table(doc, part_1_questions)

            if part_2_questions:
                self._create_part_2_answer_table(doc, part_2_questions)

            if part_3_questions:
                self._create_part_3_answer_table(doc, part_3_questions)

        except Exception as e:
            logger.error(f"Error creating THPT 2025 answer section: {e}")

    def _create_part_1_answer_table(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo bảng đáp án Phần I với logic thông minh cho format hiển thị"""
        try:
            # Tiêu đề
            section_para = doc.add_paragraph()
            section_run = section_para.add_run("PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. ")
            section_run.bold = True
            section_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}")

            note_para = doc.add_paragraph()
            note_para.add_run("(Mỗi câu trả lời đúng thí sinh được 0,25 điểm)")

            # Logic thông minh để quyết định format hiển thị đáp án
            num_questions = len(questions)
            layout_config = self._determine_answer_layout(num_questions)

            # Tạo bảng đáp án theo layout được quyết định
            self._create_smart_answer_table(doc, questions, layout_config)

        except Exception as e:
            logger.error(f"Error creating part 1 answer table: {e}")

    def _determine_answer_layout(self, num_questions: int) -> Dict[str, Any]:
        """
        Quyết định layout hiển thị đáp án dựa trên số lượng câu hỏi

        Logic:
        - Nếu đáp án ngắn (≤ 20 câu): 1 hàng
        - Nếu đáp án vừa phải (21-40 câu): 2 hàng, mỗi hàng 2 câu
        - Nếu đáp án dài (> 40 câu): 1 hàng (không auto)
        """
        if num_questions <= 20:
            # Đáp án ngắn - 1 hàng
            return {
                "type": "single_row",
                "questions_per_row": num_questions,
                "cols_per_row": num_questions + 1,  # +1 cho cột header
                "num_rows": 1
            }
        elif num_questions <= 40:
            # Đáp án vừa phải - 2 hàng, mỗi hàng tối đa 20 câu
            questions_per_row = (num_questions + 1) // 2  # Chia đều cho 2 hàng
            return {
                "type": "double_row",
                "questions_per_row": questions_per_row,
                "cols_per_row": questions_per_row + 1,
                "num_rows": 2
            }
        else:
            # Đáp án dài - 1 hàng (không auto)
            return {
                "type": "single_row_long",
                "questions_per_row": num_questions,
                "cols_per_row": min(num_questions + 1, 50),  # Giới hạn tối đa 50 cột
                "num_rows": 1
            }

    def _create_smart_answer_table(self, doc: Document, questions: List[Dict[str, Any]], layout_config: Dict[str, Any]):
        """Tạo bảng đáp án theo layout config thông minh"""
        try:
            layout_type = layout_config["type"]

            if layout_type == "single_row":
                self._create_single_row_answer_table(doc, questions, layout_config)
            elif layout_type == "double_row":
                self._create_double_row_answer_table(doc, questions, layout_config)
            elif layout_type == "single_row_long":
                self._create_single_row_long_answer_table(doc, questions, layout_config)

        except Exception as e:
            logger.error(f"Error creating smart answer table: {e}")

    def _create_single_row_answer_table(self, doc: Document, questions: List[Dict[str, Any]], layout_config: Dict[str, Any]):
        """Tạo bảng đáp án 1 hàng cho đáp án ngắn"""
        try:
            num_questions = len(questions)
            table = doc.add_table(rows=2, cols=num_questions + 1)
            table.style = 'Table Grid'

            # Header row
            table.cell(0, 0).text = "Câu"
            table.cell(1, 0).text = "Chọn"

            # Điền đáp án
            for i, question in enumerate(questions):
                table.cell(0, i + 1).text = str(i + 1)

                # Lấy đáp án đúng - thống nhất với cách lấy trong phần tạo câu hỏi
                dap_an = question.get("answer", question.get("dap_an", {}))
                correct_answer = dap_an.get("correct_answer", dap_an.get("dung", "A"))
                table.cell(1, i + 1).text = correct_answer

        except Exception as e:
            logger.error(f"Error creating single row answer table: {e}")

    def _create_double_row_answer_table(self, doc: Document, questions: List[Dict[str, Any]], layout_config: Dict[str, Any]):
        """Tạo bảng đáp án 2 hàng cho đáp án vừa phải"""
        try:
            num_questions = len(questions)
            questions_per_row = layout_config["questions_per_row"]

            table = doc.add_table(rows=4, cols=questions_per_row + 1)  # 2 hàng header + 2 hàng đáp án
            table.style = 'Table Grid'

            # Hàng 1
            table.cell(0, 0).text = "Câu"
            table.cell(1, 0).text = "Chọn"

            # Hàng 2
            table.cell(2, 0).text = "Câu"
            table.cell(3, 0).text = "Chọn"

            # Điền đáp án hàng 1
            for i in range(min(questions_per_row, num_questions)):
                table.cell(0, i + 1).text = str(i + 1)

                dap_an = questions[i].get("answer", questions[i].get("dap_an", {}))
                correct_answer = dap_an.get("correct_answer", dap_an.get("dung", "A"))
                table.cell(1, i + 1).text = correct_answer

            # Điền đáp án hàng 2 (nếu có)
            for i in range(questions_per_row, num_questions):
                col_idx = i - questions_per_row + 1
                if col_idx <= questions_per_row:
                    table.cell(2, col_idx).text = str(i + 1)

                    dap_an = questions[i].get("answer", questions[i].get("dap_an", {}))
                    correct_answer = dap_an.get("correct_answer", dap_an.get("dung", "A"))
                    table.cell(3, col_idx).text = correct_answer

        except Exception as e:
            logger.error(f"Error creating double row answer table: {e}")

    def _create_single_row_long_answer_table(self, doc: Document, questions: List[Dict[str, Any]], layout_config: Dict[str, Any]):
        """Tạo bảng đáp án 1 hàng cho đáp án dài (không auto)"""
        try:
            num_questions = len(questions)
            max_cols = layout_config["cols_per_row"]

            # Nếu quá nhiều câu, chia thành nhiều bảng
            questions_per_table = max_cols - 1  # -1 cho cột header
            num_tables = (num_questions + questions_per_table - 1) // questions_per_table

            for table_idx in range(num_tables):
                start_idx = table_idx * questions_per_table
                end_idx = min(start_idx + questions_per_table, num_questions)
                current_questions = questions[start_idx:end_idx]

                table = doc.add_table(rows=2, cols=len(current_questions) + 1)
                table.style = 'Table Grid'

                # Header row
                table.cell(0, 0).text = "Câu"
                table.cell(1, 0).text = "Chọn"

                # Điền đáp án
                for i, question in enumerate(current_questions):
                    table.cell(0, i + 1).text = str(start_idx + i + 1)

                    dap_an = question.get("answer", question.get("dap_an", {}))
                    correct_answer = dap_an.get("correct_answer", dap_an.get("dung", "A"))
                    table.cell(1, i + 1).text = correct_answer

                # Thêm khoảng cách giữa các bảng
                if table_idx < num_tables - 1:
                    doc.add_paragraph()

        except Exception as e:
            logger.error(f"Error creating single row long answer table: {e}")

    def _create_part_2_answer_table(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo bảng đáp án Phần II theo chuẩn THPT - format gộp a,b,c,d trong 1 cột"""
        try:
            doc.add_paragraph()

            section_para = doc.add_paragraph()
            section_run = section_para.add_run("PHẦN II. Câu trắc nghiệm đúng sai. ")
            section_run.bold = True
            section_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}.")
            note_para = doc.add_paragraph()
            note_para.add_run("- Thí sinh chỉ lựa chọn chính xác 01 ý trong 01 câu hỏi được 0,1 điểm;")
            note_para.add_run("\n- Thí sinh chỉ lựa chọn chính xác 02 ý trong 01 câu hỏi được 0,25 điểm;")
            note_para.add_run("\n- Thí sinh chỉ lựa chọn chính xác 03 ý trong 01 câu hỏi được 0,5 điểm;")
            note_para.add_run("\n- Thí sinh lựa chọn chính xác cả 04 ý trong 01 câu hỏi được 1 điểm.")    
            doc.add_paragraph()

            # Tạo bảng đáp án với format gộp: 1 hàng cho header, 1 hàng cho đáp án
            num_questions = len(questions)
            table = doc.add_table(rows=2, cols=num_questions + 1)
            table.style = 'Table Grid'

            # Header row
            table.cell(0, 0).text = "Câu"
            for i in range(num_questions):
                table.cell(0, i + 1).text = str(i + 1)

            # Đáp án row - gộp tất cả a,b,c,d trong 1 cell
            table.cell(1, 0).text = "Đáp án"

            for i, question in enumerate(questions):
                dap_an = question.get("answer", question.get("dap_an", {}))

                # Tạo text gộp cho a,b,c,d với đánh giá Đúng/Sai
                answer_lines = []
                for option in ["a", "b", "c", "d"]:
                    option_data = dap_an.get(option, {})

                    # Xử lý cấu trúc mới với content và evaluation
                    if isinstance(option_data, dict) and "evaluation" in option_data:
                        evaluation = option_data.get("evaluation", "Đúng")
                    else:
                        # Fallback cho cấu trúc cũ (chỉ có nội dung phát biểu)
                        evaluation = "Đúng"  # Default value

                    # Chỉ hiển thị Đúng/Sai, không hiển thị nội dung phát biểu
                    answer_lines.append(f"{option}) {evaluation}")

                # Gộp tất cả đáp án trong 1 cell
                combined_answer = "\n".join(answer_lines)
                table.cell(1, i + 1).text = combined_answer

        except Exception as e:
            logger.error(f"Error creating part 2 answer table: {e}")

    def _create_part_3_answer_table(self, doc: Document, questions: List[Dict[str, Any]]):
        """Tạo bảng đáp án Phần III"""
        try:
            doc.add_paragraph()
            
            section_para = doc.add_paragraph()
            section_run = section_para.add_run("PHẦN III. Câu trắc nghiệm trả lời ngắn. ")
            section_run.bold = True
            section_para.add_run(f"Thí sinh trả lời từ câu 1 đến câu {len(questions)}")

            note_para = doc.add_paragraph()
            note_para.add_run("Mỗi câu trả lời đúng thí sinh được 0,25 điểm.")

            # Tạo bảng đáp án
            table = doc.add_table(rows=2, cols=len(questions) + 1)
            table.style = 'Table Grid'

            # Header
            table.cell(0, 0).text = "Câu"
            table.cell(1, 0).text = "Đáp án"

            # Đáp án
            for i, question in enumerate(questions):
                table.cell(0, i + 1).text = str(i + 1)
                # Sử dụng field "answer" thay vì "dap_an"
                dap_an = question.get("answer", question.get("dap_an", {}))
                # Cho Part 3, đáp án có thể ở field "dap_an" trong answer object
                raw_answer = dap_an.get("dap_an", dap_an.get("answer", ""))
                # Đảm bảo đáp án có đúng format 4 ký tự cho phiếu tô trắc nghiệm
                formatted_answer = self._extract_numeric_answer(str(raw_answer))
                table.cell(1, i + 1).text = formatted_answer

        except Exception as e:
            logger.error(f"Error creating part 3 answer table: {e}")

    def _generate_filename(self, exam_request: Union[SmartExamRequest, Dict[str, Any]]) -> str:
        """Tạo tên file"""
        try:
            exam_title = self._get_field(exam_request, "examTitle", "Bài kiểm tra")

            # Làm sạch tên file (loại bỏ ký tự đặc biệt)
            safe_title = "".join(c for c in exam_title if c.isalnum() or c in (' ', '-', '_')).rstrip()

            # Nếu sau khi làm sạch mà rỗng thì dùng fallback
            if not safe_title.strip():
                safe_title = "Bai_kiem_tra"

            return f"{safe_title}.docx"

        except Exception as e:
            logger.error(f"Error generating filename: {e}")
            return f"Bai_kiem_tra.docx"


# Singleton instance
smart_exam_docx_service = SmartExamDocxService()
