from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
import io
from typing import Dict, Any, List
import logging
from datetime import datetime
from app.services.file_storage_service import file_storage

logger = logging.getLogger(__name__)

class DOCXService:
    """
    Service tạo file DOCX cho giáo án Hóa học
    Format theo chuẩn Bộ GD&ĐT Việt Nam
    """
    
    def __init__(self):
        # Không cần output_dir nữa vì dùng GridFS
        pass
    
    async def create_lesson_plan_docx(self, lesson_plan_data: Dict[str, Any]) -> str:
        """
        Tạo file DOCX cho giáo án và lưu vào GridFS
        
        Args:
            lesson_plan_data: Dữ liệu giáo án từ Agent
            
        Returns:
            str: GridFS file ID của DOCX đã tạo
        """
        try:
            # Tạo document mới
            doc = Document()
            
            # Thiết lập styles
            self._setup_styles(doc)
            
            # Header - Thông tin trường học
            self._add_header(doc, lesson_plan_data)
            
            # Title - Tên giáo án
            self._add_title(doc, lesson_plan_data)
            
            # Thông tin chung
            self._add_general_info(doc, lesson_plan_data)
            
            # Mục tiêu bài học
            self._add_objectives(doc, lesson_plan_data)
            
            # Nội dung bài học
            self._add_lesson_content(doc, lesson_plan_data)
            
            # Hoạt động dạy học
            self._add_teaching_activities(doc, lesson_plan_data)
            
            # Đánh giá
            self._add_assessment(doc, lesson_plan_data)
            
            # Bài tập về nhà
            self._add_homework(doc, lesson_plan_data)
            
            # Save to bytes buffer instead of file
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Generate filename
            safe_title = "".join(c for c in lesson_plan_data.get("title", "GiaoAn") 
                               if c.isalnum() or c in (' ', '-', '_')).strip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"GiaoAn_{safe_title}_{timestamp}.docx"
            
            # Store in GridFS
            metadata = {
                "lesson_plan_id": lesson_plan_data.get("id", ""),
                "lesson_title": lesson_plan_data.get("title", ""),
                "grade": lesson_plan_data.get("grade", ""),
                "topic": lesson_plan_data.get("topic", ""),
                "duration": lesson_plan_data.get("duration", 45),
                "category": "lesson_plan"
            }
            
            file_id = await file_storage.store_docx_file(
                file_content=doc_buffer.getvalue(),
                filename=filename,
                metadata=metadata
            )
            
            logger.info(f"Created and stored lesson plan DOCX: {file_id}")
            return str(file_id)
            
        except Exception as e:
            logger.error(f"Error creating lesson plan DOCX: {e}")
            raise
            
            # Mục tiêu bài học
            self._add_objectives(doc, lesson_plan_data)
            
            # Chuẩn bị
            self._add_preparation(doc, lesson_plan_data)
            
            # Tiến trình dạy học
            self._add_teaching_process(doc, lesson_plan_data)
            
            # Đánh giá
            self._add_assessment(doc, lesson_plan_data)
            
            # Bài tập về nhà
            self._add_homework(doc, lesson_plan_data)
            
            # Rút kinh nghiệm
            self._add_reflection(doc, lesson_plan_data)
            
            # Lưu file
            filename = self._generate_filename(lesson_plan_data)
            file_path = os.path.join(self.output_dir, filename)
            doc.save(file_path)
            
            logger.info(f"Created DOCX lesson plan: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to create DOCX lesson plan: {e}")
            raise
    
    def _setup_styles(self, doc: Document):
        """Thiết lập styles cho document"""
        try:
            # Style cho heading 1
            heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_font = heading1_style.font
            heading1_font.name = 'Times New Roman'
            heading1_font.size = Pt(14)
            heading1_font.bold = True
            heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading1_style.paragraph_format.space_after = Pt(12)
            
            # Style cho heading 2
            heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_font = heading2_style.font
            heading2_font.name = 'Times New Roman'
            heading2_font.size = Pt(13)
            heading2_font.bold = True
            heading2_style.paragraph_format.space_after = Pt(6)
            
            # Style cho normal text
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
            
        except Exception as e:
            logger.warning(f"Could not setup custom styles: {e}")
    
    def _add_header(self, doc: Document, data: Dict[str, Any]):
        """Thêm header thông tin trường học"""
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = header_p.add_run("TRƯỜNG THPT _______________\n")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.bold = True
        
        run2 = header_p.add_run("TỔ: HÓA HỌC")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.bold = True
        
        doc.add_paragraph()  # Khoảng trống
    
    def _add_title(self, doc: Document, data: Dict[str, Any]):
        """Thêm title giáo án"""
        title = data.get("title", f"GIÁO ÁN MÔN HÓA HỌC - LỚP {data.get('grade', '12')}")
        
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = title_p.add_run(title.upper())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Chủ đề
        topic_p = doc.add_paragraph()
        topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        topic_run = topic_p.add_run(f"Chủ đề: {data.get('topic', '')}")
        topic_run.font.name = 'Times New Roman'
        topic_run.font.size = Pt(14)
        topic_run.font.bold = True
        
        doc.add_paragraph()  # Khoảng trống
    
    def _add_general_info(self, doc: Document, data: Dict[str, Any]):
        """Thêm thông tin chung"""
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        # Dữ liệu bảng
        table_data = [
            ("Môn học:", data.get("subject", "Hóa học")),
            ("Lớp:", data.get("grade", "12")),
            ("Thời gian:", f"{data.get('duration', 45)} phút"),
            ("Loại bài học:", "Lý thuyết"),
            ("Ngày soạn:", datetime.now().strftime("%d/%m/%Y"))
        ]
        
        for i, (label, value) in enumerate(table_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            
            # Format cells
            for j in range(2):
                cell = info_table.cell(i, j)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    if j == 0:  # Label column
                        paragraph.runs[0].font.bold = True
        
        doc.add_paragraph()
    
    def _add_objectives(self, doc: Document, data: Dict[str, Any]):
        """Thêm mục tiêu bài học"""
        heading = doc.add_paragraph("I. MỤC TIÊU BÀI HỌC")
        heading.style = 'CustomHeading2'
        
        objectives = data.get("objectives", [])
        if isinstance(objectives, dict):
            # Nếu objectives là dict với các loại mục tiêu
            for obj_type, obj_list in objectives.items():
                if obj_list:
                    obj_heading = doc.add_paragraph(f"{obj_type.upper()}:")
                    obj_heading.runs[0].font.bold = True
                    obj_heading.runs[0].font.name = 'Times New Roman'
                    obj_heading.runs[0].font.size = Pt(12)
                    
                    for obj in obj_list:
                        obj_p = doc.add_paragraph(f"- {obj}")
                        obj_p.runs[0].font.name = 'Times New Roman'
                        obj_p.runs[0].font.size = Pt(12)
        else:
            # Nếu objectives là list
            for obj in objectives:
                obj_p = doc.add_paragraph(f"- {obj}")
                obj_p.runs[0].font.name = 'Times New Roman'
                obj_p.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_preparation(self, doc: Document, data: Dict[str, Any]):
        """Thêm phần chuẩn bị"""
        heading = doc.add_paragraph("II. CHUẨN BỊ")
        heading.style = 'CustomHeading2'
        
        materials = data.get("materials", [])
        
        # Giáo viên
        teacher_p = doc.add_paragraph("Giáo viên:")
        teacher_p.runs[0].font.bold = True
        teacher_p.runs[0].font.name = 'Times New Roman'
        teacher_p.runs[0].font.size = Pt(12)
        
        teacher_materials = ["Giáo án, sách giáo khoa", "Máy chiếu, bảng phụ"] + materials
        for material in teacher_materials:
            mat_p = doc.add_paragraph(f"- {material}")
            mat_p.runs[0].font.name = 'Times New Roman'
            mat_p.runs[0].font.size = Pt(12)
        
        # Học sinh
        student_p = doc.add_paragraph("Học sinh:")
        student_p.runs[0].font.bold = True
        student_p.runs[0].font.name = 'Times New Roman'
        student_p.runs[0].font.size = Pt(12)
        
        student_materials = ["Sách giáo khoa, vở ghi", "Dụng cụ học tập"]
        for material in student_materials:
            mat_p = doc.add_paragraph(f"- {material}")
            mat_p.runs[0].font.name = 'Times New Roman'
            mat_p.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_lesson_content(self, doc: Document, data: Dict[str, Any]):
        """Thêm nội dung bài học"""
        # Tiêu đề
        content_heading = doc.add_heading("III. NỘI DUNG BÀI HỌC", level=2)
        content_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Kiến thức chính
        main_content = data.get("main_content", "")
        if main_content:
            knowledge_heading = doc.add_heading("1. Kiến thức cơ bản:", level=3)
            content_para = doc.add_paragraph(main_content)
            content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Công thức và phương trình
        formulas = data.get("formulas", [])
        if formulas:
            formula_heading = doc.add_heading("2. Công thức và phương trình:", level=3)
            for i, formula in enumerate(formulas, 1):
                formula_para = doc.add_paragraph(f"{i}. {formula}")
                formula_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
        # Ví dụ minh họa
        examples = data.get("examples", [])
        if examples:
            example_heading = doc.add_heading("3. Ví dụ minh họa:", level=3)
            for i, example in enumerate(examples, 1):
                example_para = doc.add_paragraph(f"Ví dụ {i}: {example}")
                example_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_teaching_activities(self, doc: Document, data: Dict[str, Any]):
        """Thêm hoạt động dạy học"""
        # Tiêu đề
        activity_heading = doc.add_heading("IV. CÁC HOẠT ĐỘNG DẠY HỌC", level=2)
        activity_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Hoạt động khởi động
        warm_up = data.get("warm_up_activity", "")
        if warm_up:
            warmup_heading = doc.add_heading("1. Hoạt động khởi động (5 phút):", level=3)
            warmup_para = doc.add_paragraph(warm_up)
            warmup_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Hoạt động hình thành kiến thức
        main_activity = data.get("main_activity", "")
        if main_activity:
            main_heading = doc.add_heading("2. Hoạt động hình thành kiến thức (30 phút):", level=3)
            main_para = doc.add_paragraph(main_activity)
            main_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Thí nghiệm (nếu có)
        experiments = data.get("experiments", [])
        if experiments:
            exp_heading = doc.add_heading("3. Thí nghiệm:", level=3)
            for i, exp in enumerate(experiments, 1):
                exp_para = doc.add_paragraph(f"Thí nghiệm {i}: {exp.get('title', 'Không có tiêu đề')}")
                if exp.get('procedure'):
                    procedure_para = doc.add_paragraph(f"Cách tiến hành: {exp['procedure']}")
                    procedure_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    
        # Hoạt động luyện tập
        practice_activity = data.get("practice_activity", "")
        if practice_activity:
            practice_heading = doc.add_heading("4. Hoạt động luyện tập (8 phút):", level=3)
            practice_para = doc.add_paragraph(practice_activity)
            practice_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Hoạt động vận dụng
        application_activity = data.get("application_activity", "")
        if application_activity:
            application_heading = doc.add_heading("5. Hoạt động vận dụng (2 phút):", level=3)
            application_para = doc.add_paragraph(application_activity)
            application_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_assessment(self, doc: Document, data: Dict[str, Any]):
        """Thêm phần đánh giá"""
        heading = doc.add_paragraph("IV. ĐÁNH GIÁ")
        heading.style = 'CustomHeading2'
        
        assessment = data.get("assessment", {})
        
        if isinstance(assessment, dict):
            for key, value in assessment.items():
                assess_p = doc.add_paragraph(f"{key}: {value}")
                assess_p.runs[0].font.name = 'Times New Roman'
                assess_p.runs[0].font.size = Pt(12)
        else:
            assess_p = doc.add_paragraph(str(assessment))
            assess_p.runs[0].font.name = 'Times New Roman'
            assess_p.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_homework(self, doc: Document, data: Dict[str, Any]):
        """Thêm bài tập về nhà"""
        heading = doc.add_paragraph("V. BÀI TẬP VỀ NHÀ")
        heading.style = 'CustomHeading2'
        
        homework = data.get("homework", [])
        
        if homework:
            for hw in homework:
                hw_p = doc.add_paragraph(f"- {hw}")
                hw_p.runs[0].font.name = 'Times New Roman'
                hw_p.runs[0].font.size = Pt(12)
        else:
            default_hw = doc.add_paragraph("- Học thuộc bài, làm bài tập SGK")
            default_hw.runs[0].font.name = 'Times New Roman'
            default_hw.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_reflection(self, doc: Document, data: Dict[str, Any]):
        """Thêm rút kinh nghiệm"""
        heading = doc.add_paragraph("VI. RÚT KINH NGHIỆM")
        heading.style = 'CustomHeading2'
        
        reflection_p = doc.add_paragraph("(Giáo viên ghi chú sau khi dạy)")
        reflection_p.runs[0].font.italic = True
        reflection_p.runs[0].font.name = 'Times New Roman'
        reflection_p.runs[0].font.size = Pt(11)
        
        # Thêm khoảng trống để ghi chú
        for _ in range(3):
            doc.add_paragraph("_" * 80)
    
    def _generate_filename(self, data: Dict[str, Any]) -> str:
        """Tạo tên file"""
        topic = data.get("topic", "lesson")
        grade = data.get("grade", "12")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Làm sạch tên file
        safe_topic = "".join(c for c in topic if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_topic = safe_topic.replace(' ', '_')
        
        return f"GiaoAn_Hoa{grade}_{safe_topic}_{timestamp}.docx"

# Global instance
docx_service = DOCXService()
